"""
OHM - Oral History Manager
A Flet UI app designed to streamline the oral history recording and ingest workflow
for Digital.Grinnell, including WAV-to-MP3 conversion and future processing steps.
"""

import flet as ft
import os
import getpass
import logging
import csv
import json
import platform
import socket
import subprocess
import shutil
import warnings
import re
import sys
from datetime import datetime
from pathlib import Path

# Import sanitize_filename from the common-DG-utilities sibling repository
_common_dg_path = Path(__file__).resolve().parent.parent / "common-DG-utilities"
if str(_common_dg_path) not in sys.path:
    sys.path.insert(0, str(_common_dg_path))
from common_dg_utilities.dg_utils import sanitize_filename

# DOCX handling imports
try:
    from docx import Document
    from docx2pdf import convert as docx2pdf_convert
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# PDF generation imports
try:
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.lib.enums import TA_LEFT
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

# Configure logging
DATA_DIR = Path.home() / "OHM-data"
os.makedirs(DATA_DIR / "logfiles", exist_ok=True)
log_filename = DATA_DIR / "logfiles" / f"ohm_{datetime.now().strftime('%Y%m%d_%H%M%S')}.log"

file_handler = logging.FileHandler(log_filename)
file_handler.setLevel(logging.DEBUG)

console_handler = logging.StreamHandler()
console_handler.setLevel(logging.ERROR)

formatter = logging.Formatter("%(asctime)s - %(name)s - %(levelname)s - %(message)s")
file_handler.setFormatter(formatter)
console_handler.setFormatter(formatter)

logging.basicConfig(level=logging.DEBUG, handlers=[file_handler, console_handler])
logger = logging.getLogger(__name__)

# Reduce Flet's logging verbosity
logging.getLogger("flet").setLevel(logging.WARNING)
logging.getLogger("flet_core").setLevel(logging.WARNING)
logging.getLogger("flet_desktop").setLevel(logging.WARNING)

# Persistent storage file
PERSISTENCE_FILE = DATA_DIR / "persistent.json"

# Filename used when copying a permission/consent PDF into an output directory
PERMISSION_FORM_FILENAME = "permission_form.pdf"


class PersistentStorage:
    """Handle persistent storage of UI state and function usage."""

    def __init__(self):
        self.data = self.load()

    def load(self) -> dict:
        """Load persistent data from file."""
        try:
            if os.path.exists(PERSISTENCE_FILE):
                with open(PERSISTENCE_FILE, "r", encoding="utf-8") as f:
                    data = json.load(f)
                logger.info(f"Loaded persistent data from {PERSISTENCE_FILE}")
                return data
        except Exception as e:
            logger.warning(f"Could not load persistent data: {str(e)}")

        return {
            "ui_state": {
                "last_wav_dir": "",
                "last_mp3_dir": "",
                "last_input_dir": "",
                "last_output_dir": "",
                "window_left": None,
                "window_top": None,
            },
            "function_usage": {},
        }

    def save(self):
        """Save persistent data to file."""
        try:
            with open(PERSISTENCE_FILE, "w", encoding="utf-8") as f:
                json.dump(self.data, f, indent=2, ensure_ascii=False)
            logger.debug(f"Saved persistent data to {PERSISTENCE_FILE}")
        except Exception as e:
            logger.error(f"Could not save persistent data: {str(e)}")

    def set_ui_state(self, field: str, value: str):
        """Update UI state field."""
        self.data["ui_state"][field] = value
        self.save()

    def get_ui_state(self, field: str, default: str = "") -> str:
        """Get UI state field."""
        return self.data["ui_state"].get(field, default)

    def record_function_usage(self, function_name: str):
        """Record that a function was used."""
        if function_name not in self.data["function_usage"]:
            self.data["function_usage"][function_name] = {"count": 0}

        self.data["function_usage"][function_name]["last_used"] = datetime.now().isoformat()
        self.data["function_usage"][function_name]["count"] = (
            self.data["function_usage"][function_name].get("count", 0) + 1
        )
        self.save()

    def get_function_usage(self, function_name: str) -> dict:
        """Get usage stats for a function."""
        return self.data["function_usage"].get(
            function_name, {"last_used": None, "count": 0}
        )

    def get_all_function_usage(self) -> dict:
        """Get all function usage stats."""
        return self.data["function_usage"]


def check_ffmpeg() -> bool:
    """Return True if ffmpeg is available on PATH."""
    return shutil.which("ffmpeg") is not None


def convert_wav_to_mp3(
    wav_path: Path,
    mp3_path: Path,
    quality: int = 2,
    sample_rate: int = 44100,
) -> tuple[bool, str]:
    """
    Convert a WAV file to MP3 using ffmpeg.

    Args:
        wav_path: Path to the source WAV file.
        mp3_path: Destination path for the MP3 file.
        quality:  VBR quality level (0=best, 9=worst; 2 approx. 190 kbps).
        sample_rate: Output sample rate in Hz.

    Returns:
        (success, message)
    """
    if not check_ffmpeg():
        return False, (
            "ffmpeg is not installed. Please install it:\n"
            "  • macOS:  brew install ffmpeg\n"
            "  • Linux:  sudo apt install ffmpeg\n"
            "  • Windows: https://ffmpeg.org/download.html"
        )

    if not wav_path.exists():
        return False, f"Source file not found: {wav_path}"

    if mp3_path.exists():
        return False, f"Output file already exists: {mp3_path}"

    try:
        result = subprocess.run(
            [
                "ffmpeg",
                "-i", str(wav_path),
                "-codec:a", "libmp3lame",
                "-q:a", str(quality),
                "-ar", str(sample_rate),
                str(mp3_path),
                "-hide_banner",
                "-loglevel", "error",
            ],
            capture_output=True,
            text=True,
            timeout=600,
        )

        if result.returncode == 0 and mp3_path.exists():
            wav_mb = wav_path.stat().st_size / (1024 * 1024)
            mp3_mb = mp3_path.stat().st_size / (1024 * 1024)
            return True, (
                f"✅ Conversion successful!\n\n"
                f"Created: {mp3_path.name}\n"
                f"Location: {mp3_path.parent}\n\n"
                f"WAV: {wav_mb:.1f} MB  →  MP3: {mp3_mb:.1f} MB"
            )

        error_msg = result.stderr.strip() if result.stderr else "Unknown ffmpeg error"
        return False, f"❌ ffmpeg error:\n\n{error_msg}"

    except subprocess.TimeoutExpired:
        return False, "❌ Conversion timed out after 10 minutes."
    except Exception as exc:
        return False, f"❌ Unexpected error: {exc}"


def main(page: ft.Page):
    page.title = "OHM v1.2.1 - Oral History Manager"
    page.padding = 20
    page.window.width = 1050
    page.window.height = 950
    page.scroll = ft.ScrollMode.AUTO

    storage = PersistentStorage()
    logger.info("OHM application started")

    # ------------------------------------------------------------------ helpers

    def add_log_message(text: str):
        """Prepend a timestamped line to the log output field."""
        timestamp = datetime.now().strftime("%H:%M:%S")
        existing = log_output.value or ""
        log_output.value = f"[{timestamp}] {text}\n{existing}"
        page.update()

    def open_file_with_default_app(file_path: Path):
        """Open a file using the system's default application."""
        try:
            if platform.system() == "Darwin":  # macOS
                subprocess.run(["open", str(file_path)], check=True)
            elif platform.system() == "Windows":
                os.startfile(str(file_path))
            else:  # Linux and others
                subprocess.run(["xdg-open", str(file_path)], check=True)
            return True
        except Exception as ex:
            logger.error(f"Failed to open file {file_path}: {ex}")
            return False

    def update_status(message: str, is_error: bool = False, file_path: Path | None = None):
        """Update the status text widget, optionally with a clickable file link."""
        color = ft.Colors.RED_600 if is_error else ft.Colors.GREEN_700
        
        if file_path and file_path.exists():
            # Create a row with message and clickable button to open the file
            def on_open_file(e):
                if open_file_with_default_app(file_path):
                    add_log_message(f"📂 Opened file: {file_path.name}")
                else:
                    add_log_message(f"❌ Failed to open: {file_path.name}")
            
            status_container.content = ft.Row(
                controls=[
                    ft.Text(
                        message,
                        size=14,
                        color=color,
                        italic=True,
                    ),
                    ft.TextButton(
                        "Open File",
                        icon=ft.Icons.OPEN_IN_NEW,
                        on_click=on_open_file,
                        tooltip=f"Open {file_path.name}",
                    ),
                ],
                spacing=10,
                vertical_alignment=ft.CrossAxisAlignment.CENTER,
            )
        else:
            # Simple text status
            status_container.content = ft.Text(
                message,
                size=14,
                color=color,
                italic=True,
            )
        
        page.update()

    # ------------------------------------------------------------------ widgets

    # Container for status - can hold either simple text or row with button
    status_container = ft.Container(
        content=ft.Text(
            "Ready",
            size=14,
            color=ft.Colors.GREEN_700,
            italic=True,
        )
    )

    log_output = ft.TextField(
        multiline=True,
        read_only=True,
        min_lines=8,
        max_lines=12,
        text_size=12,
        border_color=ft.Colors.GREY_400,
        bgcolor=ft.Colors.GREY_100,
        value="",
    )

    # -------------------------------------------------------- Input state
    selected_file: Path | None = None
    current_directory: Path | None = None
    audio_files: list[Path] = []
    output_directory: Path | None = None  # Directory for current file's outputs
    current_epoch: int | None = None  # Epoch timestamp for current file

    # output_base_dir: the parent of all per-file output directories.
    # Mirrors <working_dir>/OHM-data and defaults to DATA_DIR until the user picks differently.
    output_base_dir: Path = DATA_DIR
    output_dir_customized: bool = False  # True once user manually picks a working directory

    # Permission PDF state
    selected_permission_pdf: Path | None = None
    pdf_files: list[Path] = []

    # Directory pickers
    directory_picker = ft.FilePicker()
    output_directory_picker = ft.FilePicker()
    page.overlay.append(directory_picker)
    page.overlay.append(output_directory_picker)

    # Input directory text field
    input_directory_field = ft.TextField(
        label="Input Directory",
        hint_text="Select a directory containing WAV/MP3 files",
        width=600,
        read_only=True,
        value=""
    )

    # Working/Output directory text field
    output_directory_field = ft.TextField(
        label="Working/Output Directory",
        hint_text="Defaults to Input Directory; OHM-data subfolder created here",
        width=600,
        read_only=True,
        value=""
    )

    # Audio file selector: read-only TextField + pick button (avoids Flutter ink layer artifacts)
    file_selection_field = ft.TextField(
        label="Select Audio File",
        hint_text="Select a directory to populate",
        width=545,
        read_only=True,
        value="",
        hint_style=ft.TextStyle(color=ft.Colors.RED_400),
    )

    pick_file_dialog_list = ft.Column(
        controls=[],
        scroll=ft.ScrollMode.AUTO,
        spacing=0,
    )

    pick_file_dialog = ft.AlertDialog(
        modal=True,
        title=ft.Text("Select Audio File"),
        content=ft.Container(
            content=pick_file_dialog_list,
            height=400,
            width=600,
        ),
        actions=[
            ft.TextButton("Cancel", on_click=lambda e: _close_pick_file_dialog(e)),
        ],
        actions_alignment=ft.MainAxisAlignment.END,
    )
    page.overlay.append(pick_file_dialog)

    # WAV-to-MP3 conversion progress dialog
    _conv_status_text = ft.Text(
        "Starting conversion…",
        size=13,
        color=ft.Colors.GREY_700,
    )
    _conv_progress = ft.ProgressBar(width=420)
    conversion_dialog = ft.AlertDialog(
        modal=True,
        title=ft.Text("Converting WAV to MP3"),
        content=ft.Container(
            content=ft.Column(
                [
                    ft.Text(
                        "⚠️  Conversion is in progress — do not close the app or "
                        "run other functions until this completes.",
                        size=13,
                        weight=ft.FontWeight.BOLD,
                    ),
                    ft.Divider(height=10, color=ft.Colors.TRANSPARENT),
                    _conv_status_text,
                    ft.Divider(height=6, color=ft.Colors.TRANSPARENT),
                    _conv_progress,
                ],
                spacing=4,
                tight=True,
            ),
            width=440,
            padding=ft.padding.only(bottom=8),
        ),
        actions=[],
        actions_alignment=ft.MainAxisAlignment.END,
    )
    page.overlay.append(conversion_dialog)

    # Permission PDF selector: read-only TextField + pick button
    pdf_selection_field = ft.TextField(
        label="Select Permission PDF",
        hint_text="Select a directory to populate",
        width=545,
        read_only=True,
        value="",
        hint_style=ft.TextStyle(color=ft.Colors.GREY_500),
    )

    pick_pdf_dialog_list = ft.Column(
        controls=[],
        scroll=ft.ScrollMode.AUTO,
        spacing=0,
    )

    pick_pdf_dialog = ft.AlertDialog(
        modal=True,
        title=ft.Text("Select Permission PDF"),
        content=ft.Container(
            content=pick_pdf_dialog_list,
            height=400,
            width=600,
        ),
        actions=[
            ft.TextButton("Cancel", on_click=lambda e: _close_pick_pdf_dialog(e)),
        ],
        actions_alignment=ft.MainAxisAlignment.END,
    )
    page.overlay.append(pick_pdf_dialog)

    # Initialize with last used input directory
    last_dir = storage.get_ui_state("last_input_dir")
    if last_dir and os.path.isdir(last_dir):
        current_directory = Path(last_dir)
        input_directory_field.value = str(current_directory)

    # Initialize working/output directory from persistence (fallback to input dir)
    last_output_dir = storage.get_ui_state("last_output_dir")
    if last_output_dir and os.path.isdir(last_output_dir):
        output_base_dir = Path(last_output_dir) / "OHM-data"
        output_directory_field.value = last_output_dir
        output_dir_customized = True
    elif current_directory:
        output_base_dir = current_directory / "OHM-data"
        output_directory_field.value = str(current_directory)

    # -------------------------------------------------------- input handlers

    def on_directory_picked(e: ft.FilePickerResultEvent):
        """Called when user selects an input directory."""
        nonlocal current_directory, audio_files, output_base_dir, selected_permission_pdf, pdf_files
        if not e.path:
            add_log_message("Directory selection cancelled")
            return

        current_directory = Path(e.path)
        input_directory_field.value = str(current_directory)
        storage.set_ui_state("last_input_dir", str(current_directory))

        # Auto-sync working/output directory when user hasn't manually customised it
        if not output_dir_customized:
            output_base_dir = current_directory / "OHM-data"
            output_directory_field.value = str(current_directory)
            os.makedirs(output_base_dir, exist_ok=True)

        # Clear file selection then auto-scan
        audio_files = []
        file_selection_field.value = ""
        file_selection_field.hint_text = "Select a directory to populate"

        # Clear permission PDF selection when directory changes
        selected_permission_pdf = None
        pdf_files = []
        pdf_selection_field.value = ""
        pdf_selection_field.hint_text = "Select a directory to populate"

        add_log_message(f"Directory selected: {current_directory}")
        update_status(f"Directory: {current_directory.name}")
        page.update()
        _scan_audio_files()
        _scan_pdf_files()

    directory_picker.on_result = on_directory_picked

    def on_output_directory_picked(e: ft.FilePickerResultEvent):
        """Called when user selects a working/output directory."""
        nonlocal output_base_dir, output_dir_customized
        if not e.path:
            add_log_message("Working/Output directory selection cancelled")
            return

        selected = Path(e.path)
        output_base_dir = selected / "OHM-data"
        os.makedirs(output_base_dir, exist_ok=True)
        output_directory_field.value = str(selected)
        storage.set_ui_state("last_output_dir", str(selected))
        output_dir_customized = True

        add_log_message(f"Working/Output Directory: {selected}")
        add_log_message(f"OHM-data folder: {output_base_dir}")
        update_status(f"Output: {selected.name}/OHM-data")
        page.update()

    output_directory_picker.on_result = on_output_directory_picked

    def on_pick_directory_click(e):
        """Open directory picker for input directory."""
        initial_dir = storage.get_ui_state("last_input_dir") or str(Path.home())
        directory_picker.get_directory_path(
            dialog_title="Select directory containing audio files",
            initial_directory=initial_dir if os.path.isdir(initial_dir) else None,
        )

    def _scan_audio_files():
        """Scan the current input directory for WAV/MP3 files and populate the dropdown."""
        nonlocal audio_files
        if not current_directory or not current_directory.exists():
            return
        try:
            audio_files = [
                f for f in current_directory.rglob("*")
                if f.is_file()
                and f.suffix.lower() in (".wav", ".mp3")
                and "Merged" not in f.relative_to(current_directory).parts
            ]
            audio_files.sort(key=lambda p: str(p.relative_to(current_directory)).lower())

            if not audio_files:
                update_status("No WAV or MP3 files found in directory or subdirectories", is_error=True)
                add_log_message(f"No audio files found in {current_directory} or its subdirectories")
                file_selection_field.value = ""
                file_selection_field.hint_text = "No audio files found in selected directory"
                page.update()
                return

            file_selection_field.hint_text = f"Select from {len(audio_files)} found audio file(s)"
            add_log_message(f"Found {len(audio_files)} audio file(s) in {current_directory.name} and subdirectories")
            update_status(f"Found {len(audio_files)} audio file(s)")
            page.update()
        except Exception as ex:
            update_status(f"Error listing files: {str(ex)}", is_error=True)
            add_log_message(f"Error listing files: {str(ex)}")

    def on_list_files_click(e):
        """List all WAV and MP3 files in the selected directory and subdirectories."""
        nonlocal audio_files
        
        if not current_directory or not current_directory.exists():
            update_status("Please select a directory first", is_error=True)
            add_log_message("No directory selected")
            return

        try:
            # Find all WAV and MP3 files recursively, excluding any Merged/ subdirectories
            audio_files = [
                f for f in current_directory.rglob("*")
                if f.is_file()
                and f.suffix.lower() in (".wav", ".mp3")
                and "Merged" not in f.relative_to(current_directory).parts
            ]
            
            # Sort by relative path
            audio_files.sort(key=lambda p: str(p.relative_to(current_directory)).lower())

            if not audio_files:
                update_status("No WAV or MP3 files found in directory or subdirectories", is_error=True)
                add_log_message(f"No audio files found in {current_directory} or its subdirectories")
                file_selection_field.value = ""
                file_selection_field.hint_text = "No audio files found in selected directory"
                page.update()
                return

            file_selection_field.hint_text = f"Select from {len(audio_files)} found audio file(s)"
            add_log_message(f"Found {len(audio_files)} audio file(s) in {current_directory.name} and subdirectories")
            update_status(f"Found {len(audio_files)} audio file(s)")
            page.update()

        except Exception as ex:
            update_status(f"Error listing files: {str(ex)}", is_error=True)
            add_log_message(f"Error listing files: {str(ex)}")

    def on_pick_output_directory_click(e):
        """Open directory picker for working/output directory."""
        initial_dir = (
            storage.get_ui_state("last_output_dir")
            or storage.get_ui_state("last_input_dir")
            or str(Path.home())
        )
        output_directory_picker.get_directory_path(
            dialog_title="Select working/output directory (OHM-data subfolder will be created here)",
            initial_directory=initial_dir if os.path.isdir(initial_dir) else None,
        )

    def _handle_file_selection(path_str):
        """Handle selection of an audio file (shared logic for dialog picker)."""
        nonlocal selected_file, output_directory, current_epoch

        if path_str:
            selected_file = Path(path_str)

            # Extract basename (filename without extension)
            basename = selected_file.stem

            # Search for existing directory with this basename
            import time
            import re
            
            # First check if the file is already in an output directory
            # Output directories match pattern: *--dg_<epoch>
            parent_dir = selected_file.parent
            if re.search(r'--dg_\d+$', parent_dir.name):
                # File is already in an output directory - reuse it
                output_directory = parent_dir
                match = re.search(r'dg_(\d+)$', parent_dir.name)
                if match:
                    current_epoch = int(match.group(1))
                    add_log_message(f"File selected: {selected_file.name}")
                    add_log_message(f"Using existing output directory: {output_directory.name}")
                    update_status(f"Selected: {selected_file.name} → {output_directory.name}")
                    logger.info(f"Reusing file's parent directory: {output_directory}")
                    return
            
            # File is not in an output directory, search for one based on basename
            existing_dirs = list(output_base_dir.glob(f"{sanitize_filename(basename).rstrip('_')}--dg_*"))
            
            if existing_dirs:
                # Reuse the first matching directory
                output_directory = existing_dirs[0]
                
                # Extract epoch from directory name using regex
                # Pattern: <basename>--dg_<epoch>
                match = re.search(r'dg_(\d+)$', output_directory.name)
                if match:
                    current_epoch = int(match.group(1))
                    add_log_message(f"File selected: {selected_file.name}")
                    add_log_message(f"Reusing existing output directory: {output_directory.name}")
                    update_status(f"Selected: {selected_file.name} → {output_directory.name}")
                    logger.info(f"Reusing output directory: {output_directory}")
                else:
                    # Shouldn't happen, but fallback to creating new
                    add_log_message("Warning: Could not extract epoch from existing directory")
                    epoch = int(time.time())
                    current_epoch = epoch
                    dirname = f"{sanitize_filename(basename).rstrip('_')}--dg_{epoch}"
                    output_directory = output_base_dir / dirname
                    try:
                        output_directory.mkdir(parents=True, exist_ok=True)
                        add_log_message(f"File selected: {selected_file.name}")
                        add_log_message(f"Created output directory: {output_directory}")
                        update_status(f"Selected: {selected_file.name} → {dirname}")
                        logger.info(f"Created output directory: {output_directory}")
                    except Exception as ex:
                        add_log_message(f"Error creating output directory: {str(ex)}")
                        update_status(f"Error creating output directory: {str(ex)}", is_error=True)
                        logger.error(f"Failed to create output directory: {str(ex)}")
                        selected_file = None
                        output_directory = None
                        current_epoch = None
            else:
                # No existing directory found, create a new one
                epoch = int(time.time())
                current_epoch = epoch
                dirname = f"{sanitize_filename(basename).rstrip('_')}--dg_{epoch}"
                output_directory = output_base_dir / dirname

                try:
                    output_directory.mkdir(parents=True, exist_ok=True)
                    add_log_message(f"File selected: {selected_file.name}")
                    add_log_message(f"Created new output directory: {output_directory}")
                    update_status(f"Selected: {selected_file.name} → {dirname}")
                    logger.info(f"Created output directory: {output_directory}")
                except Exception as ex:
                    add_log_message(f"Error creating output directory: {str(ex)}")
                    update_status(f"Error creating output directory: {str(ex)}", is_error=True)
                    logger.error(f"Failed to create output directory: {str(ex)}")
                    selected_file = None
                    output_directory = None
                    current_epoch = None
        else:
            selected_file = None
            output_directory = None
            current_epoch = None

    def _close_pick_file_dialog(e=None):
        pick_file_dialog.open = False
        page.update()

    def on_pick_file_click(e):
        """Open dialog to pick an audio file from the scanned list."""
        if not audio_files:
            update_status("No audio files found. Select a directory first.", is_error=True)
            return

        def on_file_choice(ev):
            chosen = ev.control.data
            file_selection_field.value = str(chosen.relative_to(current_directory))
            pick_file_dialog.open = False
            page.update()
            _handle_file_selection(str(chosen))
            _copy_permission_pdf_to_output()

        pick_file_dialog_list.controls = [
            ft.ListTile(
                title=ft.Text(str(f.relative_to(current_directory)), no_wrap=True),
                data=f,
                on_click=on_file_choice,
            )
            for f in audio_files
        ]
        pick_file_dialog.open = True
        page.update()

    def _scan_pdf_files():
        """Scan the current input directory for PDF files and update the PDF picker."""
        nonlocal pdf_files
        if not current_directory or not current_directory.exists():
            return
        try:
            pdf_files = sorted(
                [f for f in current_directory.rglob("*.pdf") if f.is_file()],
                key=lambda p: str(p.relative_to(current_directory)).lower(),
            )
            if pdf_files:
                pdf_selection_field.hint_text = f"Select from {len(pdf_files)} found PDF(s)"
            else:
                pdf_selection_field.hint_text = "No PDF files found in selected directory"
            page.update()
        except Exception as ex:
            add_log_message(f"Error scanning PDFs: {str(ex)}")

    def _copy_permission_pdf_to_output():
        """Copy the selected permission PDF into the current output directory."""
        if not selected_permission_pdf or not output_directory:
            return
        try:
            dest = output_directory / PERMISSION_FORM_FILENAME
            shutil.copy2(selected_permission_pdf, dest)
            add_log_message(
                f"✅ Copied permission PDF: {selected_permission_pdf.name} → {PERMISSION_FORM_FILENAME}"
            )
        except Exception as ex:
            add_log_message(f"⚠️  Could not copy permission PDF: {ex}")

    def _close_pick_pdf_dialog(e=None):
        pick_pdf_dialog.open = False
        page.update()

    def on_pick_pdf_click(e):
        """Open dialog to pick a permission PDF from the scanned list."""
        if not pdf_files:
            _scan_pdf_files()
        if not pdf_files:
            update_status("No PDF files found. Select a directory first.", is_error=True)
            return

        def on_pdf_choice(ev):
            nonlocal selected_permission_pdf
            chosen = ev.control.data
            selected_permission_pdf = chosen
            pdf_selection_field.value = chosen.name
            pick_pdf_dialog.open = False
            page.update()
            _copy_permission_pdf_to_output()

        pick_pdf_dialog_list.controls = [
            ft.ListTile(
                title=ft.Text(str(f.relative_to(current_directory)), no_wrap=True),
                data=f,
                on_click=on_pdf_choice,
            )
            for f in pdf_files
        ]
        pick_pdf_dialog.open = True
        page.update()

    # -------------------------------------------------------- function handlers

    def on_function_0_merge_audio(e):
        """Function 0: Merge two or more audio files from the current directory into one."""
        nonlocal current_directory

        if not current_directory or not current_directory.exists():
            update_status("Please select a directory first", is_error=True)
            add_log_message("No directory selected. Use Inputs section to select a directory first.")
            return

        if not check_ffmpeg():
            update_status("⚠️  ffmpeg not found — install it before merging audio files.", is_error=True)
            add_log_message("ffmpeg not installed. Install via: brew install ffmpeg (macOS)")
            return

        # Find WAV and MP3 files in current directory only (not subdirectories)
        dir_audio_files = sorted(
            [f for f in current_directory.iterdir()
             if f.is_file() and f.suffix.lower() in (".wav", ".mp3")],
            key=lambda p: p.name.lower(),
        )

        if len(dir_audio_files) < 2:
            update_status(
                f"Need at least 2 audio files in {current_directory.name} to merge.",
                is_error=True,
            )
            add_log_message(
                f"Only {len(dir_audio_files)} audio file(s) found in {current_directory.name}. "
                "Need at least 2 to merge."
            )
            return

        # Mutable selection state — list of Path objects in user-chosen order
        selected_ordered: list = []

        # ---- Dialog widgets ----
        output_name_field = ft.TextField(
            label="Output Filename",
            hint_text="Auto-detected from common name prefix",
            width=500,
            text_size=13,
        )
        merge_status_text = ft.Text("", size=13, color=ft.Colors.GREY_700, italic=True)
        selected_column = ft.Column(spacing=4, scroll=ft.ScrollMode.AUTO)
        available_column = ft.Column(spacing=4, scroll=ft.ScrollMode.AUTO)

        def compute_output_name():
            if not selected_ordered:
                return ""
            stems = [f.stem for f in selected_ordered]
            common = os.path.commonprefix(stems) if len(stems) > 1 else stems[0]
            common = common.rstrip(" -_()")
            if not common:
                common = "merged"
            ext = selected_ordered[0].suffix.lower()
            return sanitize_filename(f"{common}_MERGED{ext}")

        def refresh_dialog():
            available_column.controls.clear()
            for path in dir_audio_files:
                is_sel = path in selected_ordered
                available_column.controls.append(
                    ft.Row(
                        [
                            ft.Text(
                                path.name, size=12, expand=True,
                                color=ft.Colors.GREY_500 if is_sel else None,
                            ),
                            ft.TextButton(
                                "Remove" if is_sel else "Add →",
                                on_click=lambda ev, p=path: _toggle(p),
                                style=ft.ButtonStyle(
                                    color=ft.Colors.RED_400 if is_sel else ft.Colors.BLUE_700,
                                ),
                            ),
                        ],
                        spacing=2,
                    )
                )

            selected_column.controls.clear()
            total = len(selected_ordered)
            for idx, path in enumerate(selected_ordered):
                selected_column.controls.append(
                    ft.Row(
                        [
                            ft.Text(
                                f"{idx + 1}.",
                                width=22, size=12,
                                weight=ft.FontWeight.BOLD,
                            ),
                            ft.Text(path.name, size=12, expand=True),
                            ft.IconButton(
                                icon=ft.Icons.ARROW_UPWARD,
                                icon_size=16,
                                tooltip="Move up",
                                disabled=(idx == 0),
                                on_click=lambda ev, i=idx: _move_up(i),
                            ),
                            ft.IconButton(
                                icon=ft.Icons.ARROW_DOWNWARD,
                                icon_size=16,
                                tooltip="Move down",
                                disabled=(idx == total - 1),
                                on_click=lambda ev, i=idx: _move_down(i),
                            ),
                            ft.IconButton(
                                icon=ft.Icons.REMOVE_CIRCLE_OUTLINE,
                                icon_size=16,
                                icon_color=ft.Colors.RED_400,
                                tooltip="Remove from merge list",
                                on_click=lambda ev, i=idx: _remove(i),
                            ),
                        ],
                        spacing=2,
                    )
                )

            # Auto-update output filename when it still looks auto-generated
            auto = compute_output_name()
            current_val = output_name_field.value or ""
            if not current_val or current_val.endswith("_MERGED.wav") or current_val.endswith("_MERGED.mp3"):
                output_name_field.value = auto
            page.update()

        def _toggle(path):
            if path in selected_ordered:
                selected_ordered.remove(path)
            else:
                selected_ordered.append(path)
            refresh_dialog()

        def _move_up(idx):
            if idx > 0:
                selected_ordered[idx], selected_ordered[idx - 1] = (
                    selected_ordered[idx - 1], selected_ordered[idx]
                )
                refresh_dialog()

        def _move_down(idx):
            if idx < len(selected_ordered) - 1:
                selected_ordered[idx], selected_ordered[idx + 1] = (
                    selected_ordered[idx + 1], selected_ordered[idx]
                )
                refresh_dialog()

        def _remove(idx):
            del selected_ordered[idx]
            refresh_dialog()

        def on_merge_click(ev):
            if len(selected_ordered) < 2:
                merge_status_text.value = "⚠️  Select at least 2 files to merge."
                merge_status_text.color = ft.Colors.RED_600
                page.update()
                return

            output_filename = sanitize_filename((output_name_field.value or "").strip())
            output_name_field.value = output_filename
            if not output_filename:
                merge_status_text.value = "⚠️  Output filename cannot be empty."
                merge_status_text.color = ft.Colors.RED_600
                page.update()
                return

            output_path = current_directory / output_filename
            if output_path.exists():
                merge_status_text.value = f"⚠️  Output file already exists: {output_filename}"
                merge_status_text.color = ft.Colors.RED_600
                page.update()
                return

            exts = {f.suffix.lower() for f in selected_ordered}
            output_ext = Path(output_filename).suffix.lower()
            mixed_formats = len(exts) > 1

            merge_status_text.value = f"Merging {len(selected_ordered)} files…"
            merge_status_text.color = ft.Colors.GREY_700
            page.update()

            import tempfile
            tmp_path = None
            try:
                with tempfile.NamedTemporaryFile(
                    mode="w", suffix=".txt", delete=False, encoding="utf-8"
                ) as tmp:
                    for f in selected_ordered:
                        # Escape single quotes for ffmpeg concat format
                        path_str = str(f).replace("'", "'\\''")
                        tmp.write(f"file '{path_str}'\n")
                    tmp_path = tmp.name

                # Use stream copy when all inputs share the same container format;
                # re-encode when mixing WAV/MP3 or when the output format differs.
                if not mixed_formats and list(exts)[0] == output_ext:
                    codec_args = ["-c", "copy"]
                elif output_ext == ".mp3":
                    codec_args = ["-codec:a", "libmp3lame", "-q:a", "2", "-ar", "44100"]
                else:
                    codec_args = ["-c:a", "pcm_s16le"]

                cmd = [
                    "ffmpeg",
                    "-f", "concat",
                    "-safe", "0",
                    "-i", tmp_path,
                ] + codec_args + [
                    str(output_path),
                    "-hide_banner",
                    "-loglevel", "error",
                ]

                result = subprocess.run(cmd, capture_output=True, text=True, timeout=600)

                if result.returncode == 0 and output_path.exists():
                    size_mb = output_path.stat().st_size / (1024 * 1024)
                    merged_names = ", ".join(f.name for f in selected_ordered)
                    merge_status_text.value = (
                        f"✅ Merged: {output_filename} ({size_mb:.1f} MB)"
                    )
                    merge_status_text.color = ft.Colors.GREEN_700
                    add_log_message(
                        f"✅ Merged {len(selected_ordered)} files → {output_filename} "
                        f"({size_mb:.1f} MB)"
                    )
                    add_log_message(f"   Sources: {merged_names}")
                    add_log_message(f"   Location: {current_directory}")
                    update_status(f"✅ Merge complete: {output_filename}")
                    storage.record_function_usage("function_0_merge_audio")

                    # Write merge provenance sidecar (read back later by collect_audio_file_info)
                    now_merge = datetime.now()
                    merge_info = {
                        "merged_at": now_merge.strftime("%Y-%m-%d %H:%M:%S"),
                        "merged_at_human": now_merge.strftime("%A, %B %d, %Y at %I:%M:%S %p"),
                        "output_file": output_filename,
                        "source_count": len(selected_ordered),
                        "source_files": [
                            {
                                "order": i + 1,
                                "filename": f.name,
                                "path": str(f),
                            }
                            for i, f in enumerate(selected_ordered)
                        ],
                        "ffmpeg_codec": codec_args[1] if codec_args and len(codec_args) > 1 else "copy",
                    }
                    sidecar_path = current_directory / f"{output_path.stem}.merge_info.json"
                    try:
                        with open(sidecar_path, "w", encoding="utf-8") as _sf:
                            json.dump(merge_info, _sf, indent=2, ensure_ascii=False)
                        add_log_message(f"   Provenance: {sidecar_path.name}")
                    except Exception as _se:
                        add_log_message(f"   ⚠️  Could not write merge provenance file: {_se}")

                    # Move source files to a Merged/ subdirectory so they are excluded
                    # from future file listings and workflow statistics.
                    merged_subdir = current_directory / "Merged"
                    try:
                        merged_subdir.mkdir(exist_ok=True)
                        moved, failed = [], []
                        for src_file in selected_ordered:
                            dest = merged_subdir / src_file.name
                            # Avoid clobbering if a same-named file already exists there
                            if dest.exists():
                                stem_d, ext_d = src_file.stem, src_file.suffix
                                counter = 1
                                while dest.exists():
                                    dest = merged_subdir / f"{stem_d}_{counter}{ext_d}"
                                    counter += 1
                            try:
                                shutil.move(str(src_file), dest)
                                moved.append(src_file.name)
                            except Exception as _me:
                                failed.append(f"{src_file.name} ({_me})")
                        if moved:
                            add_log_message(
                                f"   Moved {len(moved)} source file(s) → {merged_subdir.name}/: "
                                + ", ".join(moved)
                            )
                        if failed:
                            add_log_message(
                                f"   ⚠️  Could not move: " + "; ".join(failed)
                            )
                    except Exception as _mde:
                        add_log_message(f"   ⚠️  Could not create Merged/ subdirectory: {_mde}")
                else:
                    err = (result.stderr or "").strip()
                    merge_status_text.value = f"❌ Merge failed: {err[:200]}"
                    merge_status_text.color = ft.Colors.RED_600
                    add_log_message(f"❌ Merge failed: {err}")

            except subprocess.TimeoutExpired:
                merge_status_text.value = "❌ Merge timed out (>10 minutes)."
                merge_status_text.color = ft.Colors.RED_600
                add_log_message("❌ Merge timed out after 10 minutes")
            except Exception as ex:
                merge_status_text.value = f"❌ Unexpected error: {str(ex)}"
                merge_status_text.color = ft.Colors.RED_600
                add_log_message(f"❌ Merge error: {str(ex)}")
            finally:
                if tmp_path:
                    try:
                        os.unlink(tmp_path)
                    except Exception:
                        pass

            page.update()

        def close_merge_dialog(ev):
            merge_dialog.open = False
            page.update()

        refresh_dialog()

        merge_dialog = ft.AlertDialog(
            title=ft.Text("🔀 Function 0: Merge Audio Files"),
            content=ft.Container(
                content=ft.Column(
                    [
                        ft.Text(
                            "Click 'Add →' to queue files in the desired merge order. "
                            "Use arrows to reorder. Output is saved to the same directory.",
                            size=12,
                            italic=True,
                            color=ft.Colors.GREY_700,
                        ),
                        ft.Divider(height=8),
                        ft.Row(
                            [
                                ft.Column(
                                    [
                                        ft.Text(
                                            "Available Files",
                                            size=13,
                                            weight=ft.FontWeight.BOLD,
                                        ),
                                        ft.Container(
                                            content=available_column,
                                            width=310,
                                            height=280,
                                            border=ft.border.all(1, ft.Colors.GREY_300),
                                            border_radius=4,
                                            padding=6,
                                            clip_behavior=ft.ClipBehavior.HARD_EDGE,
                                        ),
                                    ],
                                    spacing=4,
                                ),
                                ft.Container(width=16),
                                ft.Column(
                                    [
                                        ft.Text(
                                            "Merge Order (top = first)",
                                            size=13,
                                            weight=ft.FontWeight.BOLD,
                                        ),
                                        ft.Container(
                                            content=selected_column,
                                            width=360,
                                            height=280,
                                            border=ft.border.all(1, ft.Colors.GREY_300),
                                            border_radius=4,
                                            padding=6,
                                            clip_behavior=ft.ClipBehavior.HARD_EDGE,
                                        ),
                                    ],
                                    spacing=4,
                                ),
                            ],
                            spacing=0,
                            vertical_alignment=ft.CrossAxisAlignment.START,
                        ),
                        ft.Divider(height=8),
                        ft.Text(
                            "Output filename (saved in the same directory as source files):",
                            size=12,
                        ),
                        output_name_field,
                        ft.Container(height=4),
                        merge_status_text,
                    ],
                    spacing=6,
                ),
                width=730,
                height=500,
                padding=10,
            ),
            actions=[
                ft.ElevatedButton(
                    "Merge Files",
                    icon=ft.Icons.CALL_MERGE,
                    on_click=on_merge_click,
                ),
                ft.TextButton("Cancel", on_click=close_merge_dialog),
            ],
            actions_alignment=ft.MainAxisAlignment.END,
        )

        page.overlay.append(merge_dialog)
        merge_dialog.open = True
        page.update()
        add_log_message(
            f"Function 0: Merge Audio Files — {len(dir_audio_files)} files available "
            f"in {current_directory.name}"
        )
        update_status(f"Function 0: Select files to merge from {current_directory.name}")

    def on_function_1_wav_to_mp3(e):
        """Execute Function 1: WAV to MP3 Conversion"""
        import threading
        nonlocal selected_file, output_directory, current_epoch

        if not check_ffmpeg():
            update_status(
                "⚠️  ffmpeg not found — install it before converting WAV files.",
                is_error=True,
            )
            add_log_message(
                "ffmpeg not installed. Install via: brew install ffmpeg (macOS) "
                "or sudo apt install ffmpeg (Linux)"
            )
            return

        # Check if a file is selected
        if not selected_file:
            update_status("Please select an audio file first", is_error=True)
            add_log_message("No file selected. Use Inputs section to select a file.")
            return

        # Check if output directory exists
        if not output_directory or not output_directory.exists():
            update_status("Output directory not found. Please reselect the file.", is_error=True)
            add_log_message("Output directory missing. Reselect the file to recreate it.")
            return

        # Check if epoch is available
        if not current_epoch:
            update_status("Epoch timestamp not found. Please reselect the file.", is_error=True)
            add_log_message("Epoch timestamp missing. Reselect the file to regenerate it.")
            return

        # Only convert WAV files
        if selected_file.suffix.lower() != ".wav":
            update_status(
                f"Cannot convert {selected_file.suffix} file. Please select a WAV file.",
                is_error=True,
            )
            add_log_message(f"Skipped: {selected_file.name} is not a WAV file")
            return

        # Define standardized filenames using epoch
        wav_filename = sanitize_filename(f"dg_{current_epoch}.wav")
        mp3_filename = sanitize_filename(f"dg_{current_epoch}.mp3")
        wav_copy_path = output_directory / wav_filename
        mp3_path = output_directory / mp3_filename

        # Quick pre-checks that don't require I/O can stay here
        if mp3_path.exists():
            update_status(
                f"⚠️  MP3 already exists: {mp3_filename} — skipping conversion.",
                is_error=True,
            )
            add_log_message(f"Skipped: {mp3_filename} already exists in {output_directory.name}")
            return

        # Open the progress dialog NOW — before any file I/O — so it renders
        # immediately while the event handler returns and the thread does the work.
        storage.record_function_usage("function_1_wav_to_mp3")
        src_mb = selected_file.stat().st_size / (1024 * 1024)
        _conv_status_text.value = (
            f"Step 1 of 2: Copying {selected_file.name} ({src_mb:.1f} MB)…\n"
            f"Large files may take several minutes. Do not close the app."
        )
        conversion_dialog.open = True
        page.update()

        def _run_copy_and_convert():
            # --- Phase 1: copy WAV ---
            if wav_copy_path.exists():
                add_log_message(f"WAV file already exists in output directory: {wav_filename}")
            else:
                try:
                    add_log_message(f"Copying WAV file to output directory: {wav_filename}")
                    shutil.copy2(selected_file, wav_copy_path)
                    wav_size_mb = wav_copy_path.stat().st_size / (1024 * 1024)
                    add_log_message(f"✅ WAV file copied: {wav_filename} ({wav_size_mb:.1f} MB)")
                except Exception as ex:
                    conversion_dialog.open = False
                    update_status(f"Error copying WAV file: {str(ex)}", is_error=True)
                    add_log_message(f"❌ Failed to copy WAV file: {str(ex)}")
                    logger.error(f"WAV copy failed: {str(ex)}")
                    page.update()
                    return

            # --- Phase 2: convert to MP3 ---
            wav_mb = wav_copy_path.stat().st_size / (1024 * 1024)
            _conv_status_text.value = (
                f"Step 2 of 2: Converting {wav_filename} ({wav_mb:.1f} MB) → {mp3_filename}…\n"
                f"Large files may take several minutes. Do not close the app."
            )
            add_log_message(f"Starting conversion: {wav_filename} → {mp3_filename}")
            update_status(f"Converting {wav_filename} to MP3 — please wait…")
            page.update()

            success, message = convert_wav_to_mp3(wav_copy_path, mp3_path)
            conversion_dialog.open = False
            if success:
                _conv_status_text.value = "Done"
                add_log_message(f"✅ Conversion complete: {mp3_filename}")
                add_log_message(f"✅ Output location: {output_directory}")
            else:
                add_log_message(f"❌ Conversion failed: {message}")
            update_status(message.splitlines()[0], is_error=not success)
            page.update()

        threading.Thread(target=_run_copy_and_convert, daemon=True).start()


    def on_function_2_ms_word_online(e):
        """Provide instructions for transcription using MS Word Online."""
        nonlocal selected_file, output_directory, current_epoch
        
        if not selected_file:
            update_status("No file selected. Please select a file first.", is_error=True)
            add_log_message("No file selected")
            return

        # Determine which file to use for transcription
        audio_to_transcribe = None
        base_name = f"dg_{current_epoch}" if current_epoch else None
        
        # Priority: use MP3 from output directory if it exists
        if output_directory and base_name:
            mp3_in_output = output_directory / f"{base_name}.mp3"
            if mp3_in_output.exists():
                audio_to_transcribe = mp3_in_output
            elif selected_file.suffix.lower() == '.mp3':
                # Copy selected MP3 into output directory as dg_<epoch>.mp3
                try:
                    add_log_message(f"Copying MP3 to output directory as {base_name}.mp3...")
                    shutil.copy2(selected_file, mp3_in_output)
                    audio_to_transcribe = mp3_in_output
                    add_log_message(f"✅ Copied: {base_name}.mp3")
                except Exception as ex:
                    add_log_message(f"Warning: Could not copy MP3 to output directory: {str(ex)}. Using original.")
                    audio_to_transcribe = selected_file
        
        # Otherwise use selected MP3 file
        if not audio_to_transcribe and selected_file.suffix.lower() == '.mp3':
            audio_to_transcribe = selected_file
        
        if not audio_to_transcribe:
            update_status("⚠️  No MP3 file available. Use Function 1 to convert WAV to MP3 first.", is_error=True)
            add_log_message("No MP3 file found for transcription")
            return

        storage.record_function_usage("function_2_ms_word_online")
        
        # Create instructions dialog
        expected_docx_name = sanitize_filename(f"{audio_to_transcribe.stem}.docx")
        expected_json_name = sanitize_filename(f"dg_{current_epoch}_transcript.json") if current_epoch else sanitize_filename(f"{audio_to_transcribe.stem}_transcript.json")
        output_path = str(output_directory if output_directory else audio_to_transcribe.parent)
        
        # Helper function to create copyable text field
        def copyable_field(label, value):
            return ft.Column([
                ft.Text(label, size=12, weight=ft.FontWeight.BOLD),
                ft.Container(
                    content=ft.Text(
                        value,
                        size=13,
                        selectable=True,
                        color=ft.Colors.WHITE,
                    ),
                    bgcolor=ft.Colors.BLUE_900,
                    padding=8,
                    border_radius=4,
                ),
            ], spacing=3)
        
        # Show instructions in a dialog
        def close_dialog(e):
            dialog.open = False
            page.update()
        
        def on_convert_docx(e):
            """Handle DOCX to JSON conversion."""
            docx_path = Path(output_path) / expected_docx_name
            json_path = Path(output_path) / expected_json_name
            
            if not docx_path.exists():
                update_status(f"DOCX file not found: {docx_path.name}", is_error=True)
                add_log_message(f"❌ DOCX file not found: {docx_path}")
                return
            
            update_status("Converting DOCX to JSON...")
            add_log_message(f"📝 Converting {docx_path.name} to JSON...")
            
            success, message = convert_docx_to_json(
                docx_path, json_path,
                source_audio=audio_to_transcribe,
                selected_source=selected_file,
                out_dir=output_directory,
            )
            
            if success:
                update_status(f"✓ {message}")
                add_log_message(f"✓ {message}")
                add_log_message(f"  JSON: {json_path.name}")
                # Close dialog on success
                dialog.open = False
                page.update()
            else:
                update_status(message, is_error=True)
                add_log_message(f"❌ {message}")

        # ---- Review Notes (available as a tab alongside the Word instructions) ----
        notes_path = output_directory / "review_notes.md" if output_directory else None

        if notes_path and notes_path.exists():
            try:
                notes_content = notes_path.read_text(encoding="utf-8")
            except Exception:
                notes_content = ""
        else:
            notes_content = (
                f"# Review Notes\n"
                f"**File:** {audio_to_transcribe.name}  \n"
                f"**Date:** {datetime.now().strftime('%B %d, %Y')}  \n\n"
                f"## Notes\n\n"
                f"_Enter your review notes here._\n"
            )

        notes_editor = ft.TextField(
            multiline=True,
            min_lines=20,
            max_lines=30,
            value=notes_content,
            text_size=13,
            border_color=ft.Colors.GREY_400,
            bgcolor=ft.Colors.WHITE,
            expand=True,
        )

        notes_save_status = ft.Text("", size=12, italic=True, color=ft.Colors.GREY_700)

        def save_notes(ev):
            if not notes_path:
                notes_save_status.value = "❌ No output directory — notes cannot be saved."
                notes_save_status.color = ft.Colors.RED_600
                page.update()
                return
            try:
                notes_path.write_text(notes_editor.value or "", encoding="utf-8")
                notes_save_status.value = f"✅ Saved: {notes_path.name}"
                notes_save_status.color = ft.Colors.GREEN_700
                add_log_message(f"✅ Review notes saved: {notes_path}")
                update_status(f"Review notes saved for {audio_to_transcribe.name}")
                page.update()
            except Exception as ex:
                notes_save_status.value = f"❌ Save failed: {ex}"
                notes_save_status.color = ft.Colors.RED_600
                add_log_message(f"❌ Failed to save review notes: {ex}")
                page.update()

        dialog = ft.AlertDialog(
            title=ft.Text("📝 MS Word Online Transcription Instructions"),
            content=ft.Container(
                content=ft.Tabs(
                    selected_index=0,
                    animation_duration=200,
                    expand=1,
                    tabs=[
                        ft.Tab(
                            text="📝 MS Word Instructions",
                            content=ft.Column(
                                [
                                    ft.Text("Selected Audio File", size=16, weight=ft.FontWeight.BOLD),
                        copyable_field("Audio Filename:", audio_to_transcribe.name),
                        copyable_field("Audio Location:", str(audio_to_transcribe.parent)),
                        
                        ft.Divider(height=20),
                        
                        ft.Text("STEP 1: Open Microsoft Word Online", size=14, weight=ft.FontWeight.BOLD),
                        ft.Text("1. Click link below to open Word Online:"),
                        ft.TextButton(
                            "https://word.cloud.microsoft",
                            url="https://word.cloud.microsoft",
                            style=ft.ButtonStyle(color=ft.Colors.BLUE_700),
                        ),
                        ft.Text("2. Sign in with your Microsoft 365 account (subscription required)"),
                        ft.Text("3. Click 'Create Blank Document'"),
                        
                        ft.Divider(height=15),
                        
                        ft.Text("STEP 2: Set Document Name", size=14, weight=ft.FontWeight.BOLD),
                        ft.Text("1. Click on 'Document X', or whatever name is given to the new document, near the top left corner of the Word window"),
                        ft.Text("2. Replace it with this name (copy text below):"),
                        copyable_field("Document Name:", audio_to_transcribe.stem),
                        ft.Text("3. Press Enter to confirm"),
                        
                        ft.Divider(height=15),
                        
                        ft.Text("STEP 3: Start Transcription", size=14, weight=ft.FontWeight.BOLD),
                        ft.Text("1. Click the 'Home' tab (if not already selected)"),
                        ft.Text("2. Click 'Dictate' dropdown → Select 'Transcribe'"),
                        ft.Text("3. In the Transcribe pane, click 'Upload audio'"),
                        ft.Text("4. Browse to and select your audio file:"),
                        copyable_field("Audio File to Upload:", audio_to_transcribe.name),
                        ft.Text("5. Wait for transcription to complete (may take several minutes)"),
                        
                        ft.Divider(height=15),
                        
                        ft.Text("STEP 4: Review & Edit", size=14, weight=ft.FontWeight.BOLD),
                        ft.Text("1. Review the transcription in the Transcribe pane"),
                        ft.Text("2. Edit speaker names (replace 'Speaker 1' with actual names)"),
                        ft.Text("3. Fix any transcription errors"),
                        ft.Container(
                            content=ft.Text(
                                "💡  Use the Review Notes tab in this dialog to record significant changes you make to the Word-generated transcript.  Remember to SAVE your notes.",
                                size=13,
                                weight=ft.FontWeight.BOLD,
                                color=ft.Colors.BLUE_900,
                            ),
                            bgcolor=ft.Colors.LIGHT_BLUE_50,
                            border=ft.border.all(1, ft.Colors.BLUE_300),
                            border_radius=6,
                            padding=10,
                        ),
                        ft.Text(
                            "4. ⚠️  CRITICAL: Click 'Add to document' → choose 'With Speakers and Timestamps'",
                            weight=ft.FontWeight.BOLD,
                            color=ft.Colors.RED_700,
                        ),
                        ft.Text(
                            "    Choosing 'With Speakers and Timestamps' is required for speaker labels and timing to be preserved in the DOCX file.",
                            size=12,
                            italic=True,
                            color=ft.Colors.GREY_700,
                        ),
                        
                        ft.Divider(height=15),
                        
                        ft.Text("STEP 5: Save as DOCX", size=14, weight=ft.FontWeight.BOLD),
                        ft.Text("Save the DOCX directly to your output directory — no moving needed.", size=12, italic=True, color=ft.Colors.GREY_700),
                        ft.Text("Copy these values before you begin the download:", weight=ft.FontWeight.BOLD),
                        copyable_field("Save To (Output Directory):", output_path),
                        copyable_field("Save As (Filename):", expected_docx_name),
                        ft.Text("1. Click the 'File' menu in Word"),
                        ft.Text("2. Select 'Create a Copy'"),
                        ft.Text("3. Select 'Download a copy'"),
                        ft.Text("4. Click 'Download a copy' to confirm"),
                        ft.Container(
                            content=ft.Text(
                                "💡  When the browser's Save dialog appears, navigate to the Output Directory "
                                "shown above and set the filename to the Save As value above, then click Save.",
                                size=12,
                                color=ft.Colors.BLUE_900,
                            ),
                            bgcolor=ft.Colors.LIGHT_BLUE_50,
                            border=ft.border.all(1, ft.Colors.BLUE_300),
                            border_radius=6,
                            padding=10,
                        ),
                        ft.Text(
                            "If your browser saves automatically to Downloads: move the file to the Output Directory above "
                            "and rename it to match the Save As filename before clicking Convert to JSON.",
                            size=11,
                            italic=True,
                            color=ft.Colors.GREY_600,
                        ),
                        
                        ft.Divider(height=15),
                        
                        ft.Text("STEP 6: Convert to JSON", size=14, weight=ft.FontWeight.BOLD),
                        ft.Text("After moving the DOCX file to the output directory:"),
                        ft.Container(
                            content=ft.ElevatedButton(
                                "Convert to JSON",
                                icon=ft.Icons.AUTO_FIX_HIGH,
                                on_click=on_convert_docx,
                            ),
                            padding=ft.padding.only(top=10, bottom=10),
                        ),
                        ft.Text(
                            "This will automatically parse the DOCX transcription and create the JSON file.",
                            size=12,
                            italic=True,
                            color=ft.Colors.GREY_700,
                        ),
                        
                        ft.Divider(height=15),
                        
                        ft.Text("STEP 7: Generate Final Outputs", size=14, weight=ft.FontWeight.BOLD),
                        ft.Text("After conversion completes, use Function 4 to generate TXT and VTT outputs."),
                        
                        ft.Divider(height=15),
                        
                        ft.Text(
                            "Note: MS Word Online transcription is a Microsoft service. "
                            "This function provides instructions only - the actual transcription happens in your web browser.",
                            size=12,
                            italic=True,
                            color=ft.Colors.GREY_700,
                        ),
                                ],
                                scroll=ft.ScrollMode.AUTO,
                                spacing=5,
                            ),
                        ),
                        ft.Tab(
                            text="📋 Review Notes",
                            content=ft.Column(
                                [
                                    ft.Text(
                                        f"Editing: {notes_path}" if notes_path else "No output directory — notes cannot be saved.",
                                        size=11,
                                        color=ft.Colors.GREY_600,
                                        italic=True,
                                    ),
                                    ft.Container(height=6),
                                    notes_editor,
                                    ft.Container(height=4),
                                    notes_save_status,
                                    ft.Container(height=8),
                                    ft.ElevatedButton(
                                        "Save Notes",
                                        icon=ft.Icons.SAVE,
                                        on_click=save_notes,
                                    ),
                                ],
                                spacing=0,
                                scroll=ft.ScrollMode.AUTO,
                            ),
                        ),
                    ],
                ),
                width=760,
                height=680,
            ),
            actions=[
                ft.TextButton("Close", on_click=close_dialog),
            ],
            actions_alignment=ft.MainAxisAlignment.END,
        )
        
        page.overlay.append(dialog)
        dialog.open = True
        page.update()
        
        add_log_message(f"📝 Displayed MS Word Online instructions for: {audio_to_transcribe.name}")
        update_status(f"Follow the instructions to transcribe with MS Word Online")

    def convert_docx_to_json(docx_path, output_json_path, source_audio=None, selected_source=None, out_dir=None):
        """Convert MS Word transcription DOCX to JSON format.
        
        Args:
            docx_path: Path to DOCX file
            output_json_path: Path for output JSON file
            source_audio: Path of the audio file actually used for transcription
            selected_source: Path of the file the user originally selected
            out_dir: Output directory (for WAV detection)
        """
        try:
            if not DOCX_AVAILABLE:
                raise ImportError("python-docx library is required. Install it with: pip install python-docx")
            
            # Parse DOCX file
            doc = Document(docx_path)

            # Extract author metadata written by Word Online
            word_user = ""
            try:
                props = doc.core_properties
                word_user = (props.last_modified_by or props.author or "").strip()
            except Exception:
                pass

            # DEBUG: Log all paragraph content to understand the format
            add_log_message("📋 DEBUG: Parsing DOCX paragraphs...")
            paragraph_count = 0
            for p in doc.paragraphs:
                if p.text.strip():
                    paragraph_count += 1
                    if paragraph_count <= 10:  # Show first 10 non-empty paragraphs
                        add_log_message(f"  Para {paragraph_count}: {p.text[:100]}")
            add_log_message(f"  Total non-empty paragraphs: {paragraph_count}")
            
            # Extract text and parse timestamps/speakers
            # Word transcription format: timestamp and speaker on one line, text on next line(s)
            # Format: "00:00:00 Speaker Name"
            segments = []
            current_speaker = "SPEAKER_00"
            current_timestamp = None
            current_text = []
            
            # Regex pattern for Word's timestamp format: HH:MM:SS Speaker
            timestamp_speaker_pattern = re.compile(r'^(\d{1,2}):(\d{2}):(\d{2})\s+(.+)$')
            
            paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
            
            for i, text in enumerate(paragraphs):
                # Check if this line is a timestamp + speaker line
                match = timestamp_speaker_pattern.match(text)
                
                if match:
                    # Save previous segment if we have collected text
                    if current_timestamp is not None and current_text:
                        # Calculate end time (use next timestamp or add 3 seconds)
                        end_time = current_timestamp + 3.0
                        
                        segments.append({
                            "start": round(current_timestamp, 3),
                            "end": round(end_time, 3),
                            "text": " ".join(current_text),
                            "speaker": current_speaker
                        })
                    
                    # Start new segment
                    hours, minutes, seconds, docx_speaker = match.groups()
                    current_timestamp = int(hours) * 3600 + int(minutes) * 60 + int(seconds)
                    
                    # Use DOCX speaker name as-is
                    docx_speaker = docx_speaker.strip()
                    current_speaker = docx_speaker
                    current_text = []
                else:
                    # This is text content, add to current segment
                    if text and not text.startswith('Audio file') and not text.startswith('Transcript'):
                        current_text.append(text)
            
            # Don't forget the last segment
            if current_timestamp is not None and current_text:
                segments.append({
                    "start": round(current_timestamp, 3),
                    "end": round(current_timestamp + 3.0, 3),
                    "text": " ".join(current_text),
                    "speaker": current_speaker
                })
            
            # Update end times based on next segment's start time
            for i in range(len(segments) - 1):
                segments[i]["end"] = segments[i + 1]["start"]
            
            # Create JSON transcript
            notes = build_provenance_notes(
                method="MS Word Online",
                extra={
                    "ms_word_url": "https://word.cloud.microsoft",
                    "docx_source": str(docx_path),
                    **(  {"word_online_user": word_user}  if word_user else {}),
                    "segment_count": len(segments),
                    "source_audio": (
                        collect_audio_file_info(source_audio, selected_source, out_dir)
                        if source_audio and selected_source
                        else {"note": "source file info not available"}
                    ),
                    **(
                        {"permission_form": {
                            "original_filename": selected_permission_pdf.name,
                            "saved_as": PERMISSION_FORM_FILENAME,
                        }}
                        if selected_permission_pdf else {}
                    ),
                },
            )

            transcript_data = {
                "notes": notes,
                "language": "en",
                "segments": segments,
            }
            
            # Save JSON
            with open(output_json_path, 'w', encoding='utf-8') as f:
                json.dump(transcript_data, f, indent=2, ensure_ascii=False)
            
            return True, f"Successfully created JSON with {len(segments)} segments"
            
        except Exception as e:
            return False, f"Conversion failed: {str(e)}"

    def generate_pdf_from_json(json_path, pdf_path, segments):
        """Generate a formatted PDF from JSON transcript segments."""
        try:
            if not PDF_AVAILABLE:
                raise ImportError("reportlab library is required. Install it with: pip install reportlab")
            
            # Derive MP3 filename from JSON path
            json_filename = Path(json_path).name
            mp3_filename = json_filename.replace('_transcript.json', '.mp3')

            # Read narrative and speaker names from JSON notes (if present)
            narrative_text = ""
            speaker_names_for_title = []
            try:
                with open(json_path, "r", encoding="utf-8") as _f:
                    _data = json.load(_f)
                _notes = _data.get("notes", {})
                narrative_text = _notes.get("narrative", "")
                _sm = _notes.get("speaker_mapping", {})
                speaker_names_for_title = [
                    v for k, v in _sm.items()
                    if k not in ("Reviewed By", "Interviewer") and v and v.strip()
                ]
            except Exception:
                pass

            # Build document title
            if speaker_names_for_title:
                doc_title = "Oral History Transcript: " + ", ".join(speaker_names_for_title)
            else:
                doc_title = "Oral History Transcript"

            # Create PDF document
            doc = SimpleDocTemplate(
                str(pdf_path),
                pagesize=letter,
                title=doc_title,
                author="OHM — Oral History Manager",
            )
            story = []
            styles = getSampleStyleSheet()

            # Create custom styles
            title_style = ParagraphStyle(
                'DocTitle',
                parent=styles['Normal'],
                fontSize=16,
                textColor='black',
                spaceAfter=6,
                spaceBefore=0,
                alignment=TA_LEFT,
                fontName='Helvetica-Bold',
                leading=20,
            )

            narrative_style = ParagraphStyle(
                'Narrative',
                parent=styles['Normal'],
                fontSize=10,
                textColor='black',
                spaceAfter=14,
                spaceBefore=0,
                alignment=TA_LEFT,
                leading=14,
                fontName='Helvetica-Oblique',
            )

            heading_style = ParagraphStyle(
                'Heading',
                parent=styles['Normal'],
                fontSize=12,
                textColor='black',
                spaceAfter=6,
                alignment=TA_LEFT,
                fontName='Helvetica-Bold',
            )
            
            filename_style = ParagraphStyle(
                'Filename',
                parent=styles['Normal'],
                fontSize=11,
                textColor='black',
                spaceAfter=12,
                alignment=TA_LEFT,
            )
            
            timestamp_style = ParagraphStyle(
                'Timestamp',
                parent=styles['Normal'],
                fontSize=11,
                textColor='black',
                spaceAfter=4,
                alignment=TA_LEFT,
            )
            
            text_style = ParagraphStyle(
                'Text',
                parent=styles['Normal'],
                fontSize=11,
                textColor='black',
                spaceAfter=12,
                alignment=TA_LEFT,
            )

            # Document title
            story.append(Paragraph(doc_title, title_style))
            story.append(Spacer(1, 0.15 * inch))

            # Provenance narrative
            if narrative_text:
                story.append(Paragraph("Provenance", heading_style))
                story.append(Paragraph(narrative_text, narrative_style))
                story.append(Spacer(1, 0.15 * inch))
            
            # Add header section
            story.append(Paragraph("Audio file", heading_style))
            story.append(Paragraph(mp3_filename, filename_style))
            story.append(Spacer(1, 0.1 * inch))  # Single line
            story.append(Paragraph("Transcript", heading_style))
            story.append(Spacer(1, 0.1 * inch))  # Single line
            
            # Add each segment
            for segment in segments:
                start_time = segment.get('start', 0)
                speaker = segment.get('speaker', 'UNKNOWN')
                text = segment.get('text', '').strip()
                
                # Format timestamp as [HH:MM:SS]
                hours = int(start_time // 3600)
                minutes = int((start_time % 3600) // 60)
                seconds = int(start_time % 60)
                timestamp_str = f"[{hours:02d}:{minutes:02d}:{seconds:02d}] {speaker}"
                
                # Add timestamp and speaker
                story.append(Paragraph(timestamp_str, timestamp_style))
                
                # Add text
                story.append(Paragraph(text, text_style))
                
                # Double space between sections
                story.append(Spacer(1, 0.1 * inch))
            
            # Build PDF
            doc.build(story)
            return True, "PDF generated successfully"
            
        except Exception as e:
            return False, f"PDF generation failed: {str(e)}"

    def on_function_4_generate_outputs(e):
        """Generate TXT, VTT, CSV, and PDF outputs from edited JSON transcript."""
        nonlocal selected_file, output_directory, current_epoch
        
        if not selected_file:
            update_status("No file selected. Please select a file first.", is_error=True)
            add_log_message("No file selected")
            return

        if not output_directory or not output_directory.exists():
            update_status("Output directory not found. Please reselect the file.", is_error=True)
            add_log_message("Output directory missing. Reselect the file to recreate it.")
            return

        if not current_epoch:
            update_status("Epoch timestamp not found. Please reselect the file.", is_error=True)
            add_log_message("Epoch timestamp missing. Reselect the file to regenerate it.")
            return

        # Define file paths
        base_name = f"dg_{current_epoch}"
        json_path = output_directory / sanitize_filename(f"{base_name}_transcript.json")
        txt_path = output_directory / sanitize_filename(f"{base_name}.txt")
        vtt_path = output_directory / sanitize_filename(f"{base_name}.vtt")
        csv_path = output_directory / sanitize_filename(f"{base_name}.csv")
        pdf_path = output_directory / sanitize_filename(f"{base_name}.pdf")
        
        # Check if JSON exists
        if not json_path.exists():
            update_status(
                f"⚠️  Transcript JSON not found. Run Function 2 first.",
                is_error=True,
            )
            add_log_message(f"Error: {json_path.name} not found in {output_directory.name}")
            return

        storage.record_function_usage("function_4_generate_outputs")
        update_status("Generating TXT, VTT, CSV, and PDF outputs from JSON...")
        add_log_message(f"Reading transcript JSON: {json_path.name}")
        page.update()

        try:
            # Load JSON data
            with open(json_path, "r", encoding="utf-8") as f:
                data = json.load(f)
            
            segments = data.get("segments", [])
            language = data.get("language", "unknown")
            
            if not segments:
                update_status("⚠️  No segments found in JSON. Cannot generate outputs.", is_error=True)
                add_log_message("Error: JSON contains no segments")
                return
            
            # Generate TXT output with speaker labels
            add_log_message("Generating TXT output...")
            
            # Derive MP3 filename from JSON path
            mp3_filename = json_path.name.replace('_transcript.json', '.mp3')
            
            with open(txt_path, "w", encoding="utf-8") as f:
                # Add header section
                f.write("Audio file\n")
                f.write(f"{mp3_filename}\n\n")
                f.write("Transcript\n\n")
                
                # Add transcript segments with timestamps
                for segment in segments:
                    start_time = segment.get("start", 0)
                    speaker = segment.get("speaker", "UNKNOWN")
                    text = segment.get("text", "").strip()
                    
                    # Format timestamp as [HH:MM:SS]
                    hours = int(start_time // 3600)
                    minutes = int((start_time % 3600) // 60)
                    seconds = int(start_time % 60)
                    
                    f.write(f"[{hours:02d}:{minutes:02d}:{seconds:02d}] {speaker}\n")
                    f.write(f"{text}\n\n")
            
            add_log_message(f"✅ Created: {txt_path.name}")
            
            # Generate VTT output with speaker labels
            add_log_message("Generating VTT output...")
            with open(vtt_path, "w", encoding="utf-8") as f:
                f.write("WEBVTT\n\n")
                for segment in segments:
                    start_time = format_vtt_timestamp(segment.get("start", 0))
                    end_time = format_vtt_timestamp(segment.get("end", 0))
                    speaker = segment.get("speaker", "UNKNOWN")
                    text = segment.get("text", "").strip()
                    
                    f.write(f"{start_time} --> {end_time}\n")
                    f.write(f"<v {speaker}>{text}</v>\n\n")
            
            add_log_message(f"✅ Created: {vtt_path.name}")

            # Generate CSV output
            add_log_message("Generating CSV output...")
            with open(csv_path, "w", newline="", encoding="utf-8") as f:
                writer = csv.writer(f)
                writer.writerow(["timestamp", "speaker", "words"])
                for segment in segments:
                    start_time = segment.get("start", 0)
                    hours = int(start_time // 3600)
                    minutes = int((start_time % 3600) // 60)
                    seconds = int(start_time % 60)
                    timestamp = f"{hours:02d}:{minutes:02d}:{seconds:02d}"
                    speaker = segment.get("speaker", "UNKNOWN")
                    words = segment.get("text", "").strip()
                    writer.writerow([timestamp, speaker, words])
            add_log_message(f"✅ Created: {csv_path.name}")

            # Generate PDF output
            add_log_message("Generating PDF output...")
            pdf_success, pdf_message = generate_pdf_from_json(json_path, pdf_path, segments)
            if pdf_success:
                add_log_message(f"✅ Created: {pdf_path.name}")
            else:
                add_log_message(f"⚠️  PDF generation warning: {pdf_message}")
            
            add_log_message(f"✅ Output generation complete! Language: {language}")
            add_log_message(f"✅ Output location: {output_directory}")
            
            success_msg = f"✅ Generated TXT, VTT, CSV, and PDF outputs from edited JSON!"
            update_status(success_msg)
            
        except json.JSONDecodeError as ex:
            error_msg = f"❌ Invalid JSON format: {str(ex)}"
            add_log_message(error_msg)
            update_status(error_msg, is_error=True)
            logger.error(f"JSON decode error: {str(ex)}")
        except Exception as ex:
            error_msg = f"❌ Output generation failed: {str(ex)}"
            add_log_message(error_msg)
            update_status(error_msg, is_error=True)
            logger.error(f"Output generation error: {str(ex)}")
        
        page.update()

    def format_vtt_timestamp(seconds):
        """Convert seconds to VTT timestamp format (HH:MM:SS.mmm)"""
        hours = int(seconds // 3600)
        minutes = int((seconds % 3600) // 60)
        secs = int(seconds % 60)
        millis = int((seconds % 1) * 1000)
        return f"{hours:02d}:{minutes:02d}:{secs:02d}.{millis:03d}"

    def on_function_5_report_progress(e):
        """Generate a workflow progress report comparing input directory with processed files."""
        nonlocal current_directory
        
        if not current_directory or not current_directory.exists():
            update_status("⚠️  No input directory selected. Please select a directory first.", is_error=True)
            add_log_message("No input directory selected")
            return
        
        storage.record_function_usage("function_5_report_progress")
        update_status("Generating workflow progress report...")
        add_log_message("Scanning input directory and OHM-data...")
        page.update()
        
        try:
            # Scan input directory for audio files, excluding any Merged/ subdirectories
            input_files = {}
            for file_path in current_directory.rglob("*"):
                if (
                    file_path.is_file()
                    and file_path.suffix.lower() in (".wav", ".mp3")
                    and "Merged" not in file_path.relative_to(current_directory).parts
                ):
                    stem = file_path.stem
                    if stem not in input_files:
                        input_files[stem] = {'input_path': file_path, 'format': file_path.suffix.lower()}
            
            # Scan OHM-data for processed files
            processed_files = {}
            if output_base_dir.exists():
                for dir_path in output_base_dir.iterdir():
                    if dir_path.is_dir() and '--dg_' in dir_path.name:
                        # Extract basename and epoch from directory name
                        parts = dir_path.name.split('--dg_')
                        if len(parts) == 2:
                            original_basename = parts[0]  # e.g., "Darrell Hall"
                            epoch = parts[1]
                            dg_name = f"dg_{epoch}"
                            
                            processed_files[original_basename] = {
                                'directory': dir_path,
                                'dg_name': dg_name,
                                'wav': (dir_path / f"{dg_name}.wav").exists(),
                                'mp3': (dir_path / f"{dg_name}.mp3").exists(),
                                'json': (dir_path / f"{dg_name}_transcript.json").exists(),
                                'txt': (dir_path / f"{dg_name}.txt").exists(),
                                'vtt': (dir_path / f"{dg_name}.vtt").exists(),
                                'csv': (dir_path / f"{dg_name}.csv").exists(),
                                'pdf': (dir_path / f"{dg_name}.pdf").exists(),
                                'notes': (dir_path / "review_notes.md").exists(),
                            }
            
            # Calculate statistics
            total_input = len(input_files)
            total_processed = len(processed_files)
            
            complete_count = sum(1 for f in processed_files.values() 
                               if f['mp3'] and f['json'] and f['txt'] and f['vtt'] and f['csv'] and f['pdf'])
            in_progress_count = sum(1 for f in processed_files.values()
                                  if not (f['mp3'] and f['json'] and f['txt'] and f['vtt'] and f['csv'] and f['pdf'])
                                  and (f['mp3'] or f['json'] or f['txt']))
            
            mp3_count = sum(1 for f in processed_files.values() if f['mp3'])
            json_count = sum(1 for f in processed_files.values() if f['json'])
            txt_count = sum(1 for f in processed_files.values() if f['txt'])
            vtt_count = sum(1 for f in processed_files.values() if f['vtt'])
            csv_count = sum(1 for f in processed_files.values() if f['csv'])
            pdf_count = sum(1 for f in processed_files.values() if f['pdf'])
            notes_count = sum(1 for f in processed_files.values() if f['notes'])

            # Compute unmatched input files using sanitized stems (dirs use sanitize_filename)
            unprocessed = [name for name in sorted(input_files.keys())
                           if sanitize_filename(name).rstrip('_') not in processed_files]

            # Generate report content
            timestamp = datetime.now()
            report_content = f"""# OHM Workflow Progress Report

**Generated:** {timestamp.strftime('%B %d, %Y at %I:%M %p')}  
**Auto-generated by OHM App - Function 5**

## Overview

This report tracks the processing status of audio files from the input directory through the OHM workflow.

**Input Directory:** `{current_directory}`  
**Output Directory:** `{output_base_dir}`

## Workflow Stages

1. **Audio (WAV/MP3)** - Original or converted audio
2. **Transcription (JSON)** - Transcript from Function 2 (MS Word Online)
3. **Outputs (TXT, VTT, CSV, PDF)** - Final deliverables from Function 4

## Progress Legend

- ✅ **Complete** - All stages finished (MP3, JSON, TXT, VTT, CSV, PDF)
- 🟡 **In Progress** - Some stages completed
- ⏳ **Not Started** - Only source file exists in input directory

---

## Summary Statistics

| Metric | Count |
|--------|-------|
| Files in Input Directory | {total_input} |
| Processed Directories | {total_processed} |
| ✅ Complete (All stages) | {complete_count} |
| 🟡 In Progress (Some stages) | {in_progress_count} |
| ⏳ Not Started | {len(unprocessed)} |

### Processing Stages Completed

| Stage | Count | Percentage |
|-------|-------|------------|
| MP3 Files | {mp3_count}/{total_processed if total_processed > 0 else 1} | {mp3_count/total_processed*100 if total_processed > 0 else 0:.0f}% |
| JSON Transcripts | {json_count}/{total_processed if total_processed > 0 else 1} | {json_count/total_processed*100 if total_processed > 0 else 0:.0f}% |
| TXT Outputs | {txt_count}/{total_processed if total_processed > 0 else 1} | {txt_count/total_processed*100 if total_processed > 0 else 0:.0f}% |
| VTT Outputs | {vtt_count}/{total_processed if total_processed > 0 else 1} | {vtt_count/total_processed*100 if total_processed > 0 else 0:.0f}% |
| CSV Outputs | {csv_count}/{total_processed if total_processed > 0 else 1} | {csv_count/total_processed*100 if total_processed > 0 else 0:.0f}% |
| PDF Outputs | {pdf_count}/{total_processed if total_processed > 0 else 1} | {pdf_count/total_processed*100 if total_processed > 0 else 0:.0f}% |
| Review Notes | {notes_count}/{total_processed if total_processed > 0 else 1} | {notes_count/total_processed*100 if total_processed > 0 else 0:.0f}% |

---

## Processed Files Status

"""
            
            # Add complete files
            complete_files = [(name, info) for name, info in sorted(processed_files.items())
                            if info['mp3'] and info['json'] and info['txt'] and info['vtt'] and info['csv'] and info['pdf']]
            
            if complete_files:
                report_content += f"### ✅ Complete ({len(complete_files)} files)\n\n"
                for name, info in complete_files:
                    notes_indicator = "✅" if info['notes'] else "—"
                    report_content += f"**{name}** (`{info['dg_name']}`)  \n"
                    report_content += f"- ✅ MP3, JSON, TXT, VTT, CSV, PDF\n"
                    report_content += f"- Review Notes: {notes_indicator}\n"
                    report_content += f"- Location: `{info['directory'].name}`\n\n"
            
            # Add in-progress files
            in_progress_files = [(name, info) for name, info in sorted(processed_files.items())
                               if not (info['mp3'] and info['json'] and info['txt'] and info['vtt'] and info['csv'] and info['pdf'])
                               and (info['mp3'] or info['json'] or info['txt'])]
            
            if in_progress_files:
                report_content += f"### 🟡 In Progress ({len(in_progress_files)} files)\n\n"
                for name, info in in_progress_files:
                    missing = []
                    if not info['mp3']: missing.append('MP3')
                    if not info['json']: missing.append('JSON')
                    if not info['txt']: missing.append('TXT')
                    if not info['vtt']: missing.append('VTT')
                    if not info['csv']: missing.append('CSV')
                    if not info['pdf']: missing.append('PDF')
                    
                    notes_indicator = "✅" if info['notes'] else "—"
                    report_content += f"**{name}** (`{info['dg_name']}`)  \n"
                    report_content += f"- Missing: {', '.join(missing)}\n"
                    report_content += f"- Review Notes: {notes_indicator}\n"
                    report_content += f"- Location: `{info['directory'].name}`\n\n"
            
            # Add unprocessed files (already computed above with sanitized matching)
            
            if unprocessed:
                report_content += f"### ⏳ Not Started ({len(unprocessed)} files)\n\n"
                for name in unprocessed:
                    file_info = input_files[name]
                    report_content += f"**{name}{file_info['format']}**  \n"
                    report_content += f"- Status: Source file only\n"
                    report_content += f"- Next: Run Function 1 (if WAV) or Function 2 (if MP3)\n\n"
            
            report_content += f"""---

## Workflow Reminder

For each audio file:

1. **(If WAV) Function 1: Convert WAV to MP3**
2. **Function 2: Transcribe** (MS Word Online mode)
3. **Edit JSON** (fix speaker names, correct text)
4. **Function 4: Generate TXT, VTT, CSV & PDF** from edited JSON

---

**Next Report:** Run Function 5 again to update this status report
"""
            
            # Save report with timestamp
            report_filename = sanitize_filename(f"workflow_progress_{timestamp.strftime('%Y%m%d_%H%M%S')}.md")
            report_path = output_base_dir / report_filename
            
            with open(report_path, 'w', encoding='utf-8') as f:
                f.write(report_content)
            
            success_msg = f"✅ Progress report generated: {report_filename}"
            add_log_message(success_msg)
            add_log_message(f"  Saved to: {output_base_dir}")
            add_log_message(f"  Found {total_input} input files, {total_processed} processed")
            update_status(success_msg, file_path=report_path)
            
        except Exception as ex:
            error_msg = f"❌ Report generation failed: {str(ex)}"
            add_log_message(error_msg)
            update_status(error_msg, is_error=True)
            logger.error(f"Report generation error: {str(ex)}")
        
        page.update()

    def on_function_3_review_notes(e):
        """Function 3: Create or edit a review_notes.md file for the selected oral history."""
        nonlocal selected_file, output_directory

        if not selected_file:
            update_status("Please select an audio file first", is_error=True)
            add_log_message("No file selected. Use Inputs section to select a file first.")
            return

        if not output_directory or not output_directory.exists():
            update_status("Output directory not found. Please reselect the file.", is_error=True)
            add_log_message("Output directory missing. Reselect the file to recreate it.")
            return

        notes_path = output_directory / "review_notes.md"

        # Load existing content or seed with a template
        if notes_path.exists():
            try:
                existing_content = notes_path.read_text(encoding="utf-8")
            except Exception as ex:
                existing_content = ""
                add_log_message(f"⚠️  Could not read existing review notes: {ex}")
        else:
            existing_content = (
                f"# Review Notes\n"
                f"**File:** {selected_file.name}  \n"
                f"**Date:** {datetime.now().strftime('%B %d, %Y')}  \n\n"
                f"## Notes\n\n"
                f"_Enter your review notes here._\n"
            )

        storage.record_function_usage("function_3_review_notes")

        # Editor widget
        editor = ft.TextField(
            multiline=True,
            min_lines=20,
            max_lines=30,
            value=existing_content,
            text_size=13,
            border_color=ft.Colors.GREY_400,
            bgcolor=ft.Colors.WHITE,
            expand=True,
        )

        save_status = ft.Text("", size=12, italic=True, color=ft.Colors.GREY_700)

        def do_save(ev):
            try:
                notes_path.write_text(editor.value or "", encoding="utf-8")
                save_status.value = f"✅ Saved: {notes_path.name}"
                save_status.color = ft.Colors.GREEN_700
                add_log_message(f"✅ Review notes saved: {notes_path}")
                update_status(f"Review notes saved for {selected_file.name}")
                page.update()
            except Exception as ex:
                save_status.value = f"❌ Save failed: {ex}"
                save_status.color = ft.Colors.RED_600
                add_log_message(f"❌ Failed to save review notes: {ex}")
                page.update()

        def do_save_close(ev):
            do_save(ev)
            if "❌" not in (save_status.value or ""):
                notes_dialog.open = False
                page.update()

        def do_cancel(ev):
            notes_dialog.open = False
            page.update()

        notes_dialog = ft.AlertDialog(
            modal=True,
            title=ft.Text(
                f"📋 Review Notes — {selected_file.name}",
                weight=ft.FontWeight.BOLD,
            ),
            content=ft.Container(
                content=ft.Column(
                    [
                        ft.Text(
                            f"Editing: {notes_path}",
                            size=11,
                            color=ft.Colors.GREY_600,
                            italic=True,
                        ),
                        ft.Container(height=6),
                        editor,
                        ft.Container(height=4),
                        save_status,
                    ],
                    spacing=0,
                    tight=True,
                ),
                width=820,
                height=560,
                padding=10,
            ),
            actions=[
                ft.ElevatedButton(
                    "Save",
                    icon=ft.Icons.SAVE,
                    on_click=do_save_close,
                ),
                ft.TextButton("Cancel", on_click=do_cancel),
            ],
            actions_alignment=ft.MainAxisAlignment.END,
        )

        page.overlay.append(notes_dialog)
        notes_dialog.open = True
        page.update()
        add_log_message(f"📋 Opened review notes editor for {selected_file.name}")
        update_status(f"Editing review notes for {selected_file.name}")

    def on_placeholder_function(e):
        """Placeholder for future functions"""
        add_log_message("This function is not yet implemented")
        update_status("Placeholder function — not yet implemented", is_error=False)

    def on_copy_status_click(e):
        """Copy current status text to clipboard."""
        # Extract text from status_container content
        status_text_value = ""
        if isinstance(status_container.content, ft.Text):
            status_text_value = status_container.content.value or ""
        elif isinstance(status_container.content, ft.Row):
            # Get text from first control (Text widget)
            if status_container.content.controls:
                first_control = status_container.content.controls[0]
                if isinstance(first_control, ft.Text):
                    status_text_value = first_control.value or ""
        
        page.set_clipboard(status_text_value)
        add_log_message("Status copied to clipboard.")

    def on_clear_log_click(e):
        """Clear the log output field."""
        log_output.value = ""
        page.update()

    # -------------------------------------------------------- function metadata

    def on_function_2_transcribe(e):
        """Transcription function — uses MS Word Online."""
        on_function_2_ms_word_online(e)

    # Active functions - frequently used
    active_functions = [
        "function_0_merge_audio",
        "function_1_wav_to_mp3",
        "function_2_transcribe",
        "function_3_review_notes",
        "function_4_generate_outputs",
        "function_5_report_progress",
    ]

    functions = {
        "function_0_merge_audio": {
            "label": "0: Merge Audio Files",
            "icon": "🔀",
            "handler": on_function_0_merge_audio,
            "help_file": "FUNCTION_0_MERGE_AUDIO.md"
        },
        "function_1_wav_to_mp3": {
            "label": "1: Convert WAV to MP3",
            "icon": "🎵",
            "handler": on_function_1_wav_to_mp3,
            "help_file": "FUNCTION_1_WAV_TO_MP3.md"
        },
        "function_2_transcribe": {
            "label": "2: Transcribe with MS Word Online",
            "icon": "📝",
            "handler": on_function_2_transcribe,
            "help_file": "FUNCTION_2_MS_WORD_ONLINE.md"
        },
        "function_3_review_notes": {
            "label": "3: Edit Review Notes",
            "icon": "📋",
            "handler": on_function_3_review_notes,
            "help_file": "FUNCTION_3_REVIEW_NOTES.md"
        },
        "function_4_generate_outputs": {
            "label": "4: Generate TXT, VTT, CSV & PDF from JSON",
            "icon": "📄",
            "handler": on_function_4_generate_outputs,
            "help_file": "FUNCTION_4_GENERATE_OUTPUTS.md"
        },
        "function_5_report_progress": {
            "label": "5: Report Workflow Progress",
            "icon": "📊",
            "handler": on_function_5_report_progress,
            "help_file": "FUNCTION_5_REPORT_PROGRESS.md"
        },

    }

    # Help Mode checkbox state
    help_mode_enabled = ft.Ref[ft.Checkbox]()

    def show_help_dialog(function_key):
        """Display the help markdown file for a function"""
        if function_key not in functions:
            return

        func_info = functions[function_key]
        help_file = func_info.get("help_file")
        display_label = func_info['label']  # Default to function label

        if not help_file:
            add_log_message(f"No help file available for {display_label}")
            return

        try:
            # Read the markdown file
            with open(help_file, "r", encoding="utf-8") as f:
                markdown_content = f.read()

            add_log_message(f"Displaying help for: {display_label}")

            def close_help_dialog(e):
                help_dialog.open = False
                page.update()

            def copy_help(e):
                page.set_clipboard(markdown_content)
                copy_help_button.text = "Copied!"
                page.update()
                # Reset button text after 2 seconds
                import threading

                def reset_text():
                    import time
                    time.sleep(2)
                    copy_help_button.text = "Copy to Clipboard"
                    page.update()

                threading.Thread(target=reset_text, daemon=True).start()

            copy_help_button = ft.TextButton("Copy to Clipboard", on_click=copy_help)

            help_dialog = ft.AlertDialog(
                modal=True,
                title=ft.Text(f"Function {display_label}", weight=ft.FontWeight.BOLD),
                content=ft.Container(
                    content=ft.Column(
                        [
                            ft.Text(
                                f"File: {help_file}",
                                size=11,
                                color=ft.Colors.GREY_600,
                                italic=True,
                            ),
                            ft.Container(height=10),
                            ft.Container(
                                content=ft.Column(
                                    [
                                        ft.Markdown(
                                            value=markdown_content,
                                            selectable=True,
                                            extension_set=ft.MarkdownExtensionSet.GITHUB_WEB,
                                            on_tap_link=lambda e: page.launch_url(e.data),
                                        ),
                                    ],
                                    scroll=ft.ScrollMode.AUTO,
                                ),
                                width=800,
                                height=600,
                                padding=10,
                            ),
                        ],
                        tight=True,
                    ),
                    padding=10,
                ),
                actions=[
                    copy_help_button,
                    ft.TextButton("Close", on_click=close_help_dialog),
                ],
                actions_alignment=ft.MainAxisAlignment.END,
            )

            page.overlay.append(help_dialog)
            help_dialog.open = True
            page.update()

        except FileNotFoundError:
            add_log_message(f"Help file not found: {help_file}")
            update_status(f"Help file not found: {help_file}", True)
        except Exception as e:
            add_log_message(f"Error reading help file: {str(e)}")
            update_status(f"Error reading help file: {str(e)}", True)

    def build_provenance_notes(method: str, extra: dict | None = None) -> dict:
        """Build a provenance/notes dict capturing how and when a transcript was created."""
        import sys
        now = datetime.now()

        # System / machine info
        try:
            hostname = socket.gethostname()
        except Exception:
            hostname = "unknown"
        try:
            fqdn = socket.getfqdn()
        except Exception:
            fqdn = hostname

        notes = {
            "created_at": now.strftime("%Y-%m-%d %H:%M:%S"),
            "created_at_human": now.strftime("%A, %B %d, %Y at %I:%M:%S %p"),
            "transcription_method": method,
            "app": "OHM — Oral History Manager",
            "system": {
                "hostname": hostname,
                "fqdn": fqdn,
                "os": platform.platform(),
                "os_name": platform.system(),
                "os_version": platform.version(),
                "machine": platform.machine(),
                "processor": platform.processor(),
                "python_version": sys.version,
            },
        }

        if extra:
            notes.update(extra)

        # Build human-readable narrative paragraph
        try:
            audio_info = notes.get("source_audio", {})
            tech = audio_info.get("audio_technical", {})
            sm = notes.get("speaker_mapping", {})

            # Date / method sentence
            narrative = (
                f"This transcript was created on {notes['created_at_human']} "
                f"using the {notes['app']} application."
            )

            # Transcription method
            narrative += f" The transcription method was \"{notes['transcription_method']}\"."

            # MS Word URL
            if "ms_word_url" in notes:
                narrative += f" The audio was transcribed via Microsoft Word Online ({notes['ms_word_url']})."

            # Language / segments
            if "detected_language" in notes:
                narrative += f" The detected language was \"{notes['detected_language']}\"."
            segs = notes.get("segment_count")
            if segs is not None:
                narrative += f" The transcript contains {segs} segments."

            # Source audio
            src_file = audio_info.get("selected_file", {}).get("filename", "")
            xcd_file = audio_info.get("transcribed_file", {}).get("filename", "")
            started_wav = audio_info.get("started_from_wav")
            if src_file:
                if started_wav:
                    narrative += (
                        f" The source audio originated as a WAV file (\"{src_file}\"),"
                        f" which was converted to MP3 (\"{xcd_file}\") before transcription."
                    )
                else:
                    narrative += f" The source audio file was \"{src_file}\"."
            wav_list = audio_info.get("wav_in_output_directory", [])
            if wav_list:
                wav_names = ", ".join(Path(w).name for w in wav_list)
                narrative += f" A WAV file ({wav_names}) is also present in the output directory."

            # Merge provenance
            merge_info = audio_info.get("merge_info")
            if merge_info:
                sources = merge_info.get("source_files", [])
                source_names = ", ".join(f'"{s["filename"]}"' for s in sources)
                merged_at_str = merge_info.get("merged_at_human", merge_info.get("merged_at", ""))
                narrative += (
                    f" This audio file was created by merging {len(sources)} source file(s)"
                    f" ({source_names}) on {merged_at_str} using Function 0"
                    f" (Merge Audio Files) of the OHM application."
                )

            # Technical audio details
            if tech and "error" not in tech:
                dur = tech.get("duration_human", "")
                codec = tech.get("codec", "")
                sr = tech.get("sample_rate_hz", "")
                br = tech.get("bit_rate_kbps", "")
                parts = []
                if dur:
                    parts.append(f"duration {dur}")
                if codec:
                    parts.append(f"codec {codec}")
                if sr:
                    parts.append(f"sample rate {sr} Hz")
                if br:
                    parts.append(f"bit rate {br} kbps")
                if parts:
                    narrative += f" Audio technical details: {', '.join(parts)}."

            # Speaker mapping
            interviewer_val = sm.get("Interviewer", "")
            interviewer_named = bool(
                interviewer_val and interviewer_val.strip()
                and interviewer_val.strip().lower() != "interviewer"
            )
            # Only list Speaker N entries with real values
            named_speakers = {
                k: v for k, v in sm.items()
                if k not in ("Reviewed By", "Interviewer") and v and v.strip()
            }
            # Include Interviewer only when a real name was provided
            if interviewer_named:
                named_speakers = {"Interviewer": interviewer_val.strip(), **named_speakers}
            reviewer = sm.get("Reviewed By", "")
            if named_speakers:
                speaker_list = "; ".join(f"{k}: {v}" for k, v in named_speakers.items())
                narrative += f" The individual(s) identified as speaker(s) are: {speaker_list}."
            if not interviewer_named:
                narrative += " Interviewer was not specifically identified."
            if reviewer:
                narrative += f" The transcript was reviewed and edited by {reviewer}."

            # System / machine — omit hostname for privacy; show machine type and OS only
            sys_info = notes.get("system", {})
            os_name = sys_info.get("os_name", "")
            machine = sys_info.get("machine", "")
            os_display = {"Darwin": "macOS", "Windows": "Windows", "Linux": "Linux"}.get(os_name, os_name)
            machine_display = "Apple Silicon Mac" if machine in ("arm64", "aarch64") else (
                "Intel Mac" if os_name == "Darwin" else (
                    "PC" if os_name in ("Windows", "Linux") else machine
                )
            )
            if os_display or machine_display:
                parts = [p for p in (machine_display, os_display) if p]
                narrative += f" Processing was performed on a {', '.join(parts)} system."

            # OS user who ran the app
            os_user = sys_info.get("os_user", "")
            if os_user:
                narrative += f" The app was run by OS user \u201c{os_user}\u201d."

            # MS Word Online reviewer (from DOCX core properties, if captured)
            word_user = notes.get("word_online_user", "")
            if word_user:
                narrative += f" The Word Online document was last modified by \u201c{word_user}\u201d."

            # Permission / consent form
            perm_form = notes.get("permission_form")
            if perm_form:
                orig_name = perm_form.get("original_filename", "")
                saved_as = perm_form.get("saved_as", PERMISSION_FORM_FILENAME)
                if orig_name:
                    narrative += (
                        f" \"{orig_name}\" was specified as the permissions form for this"
                        f" object and is saved as '{saved_as}'."
                    )

            notes["narrative"] = narrative.strip()
        except Exception:
            notes["narrative"] = "Narrative generation failed."

        # Reorder so narrative is always the first key in the notes dict
        return {"narrative": notes.pop("narrative"), **notes}

    def collect_audio_file_info(audio_path: Path, selected_path: Path, out_dir: Path | None = None) -> dict:
        """Collect technical metadata about the source audio file(s)."""
        info: dict = {}

        def file_stats(p: Path) -> dict:
            try:
                st = p.stat()
                return {
                    "path": str(p),
                    "filename": p.name,
                    "size_bytes": st.st_size,
                    "size_human": f"{st.st_size / (1024 * 1024):.2f} MB",
                    "modified": datetime.fromtimestamp(st.st_mtime).strftime("%Y-%m-%d %H:%M:%S"),
                }
            except Exception as exc:
                return {"path": str(p), "error": str(exc)}

        # The file actually transcribed
        info["transcribed_file"] = file_stats(audio_path)
        info["transcribed_file_type"] = audio_path.suffix.lower().lstrip(".")

        # The file the user originally selected
        info["selected_file"] = file_stats(selected_path)
        info["selected_file_type"] = selected_path.suffix.lower().lstrip(".")
        info["started_from_wav"] = selected_path.suffix.lower() == ".wav"

        # Check for WAV counterpart in output directory
        if out_dir:
            wav_files = list(out_dir.glob("*.wav")) + list(out_dir.glob("*.WAV"))
            info["wav_in_output_directory"] = [str(w) for w in wav_files] if wav_files else []
        else:
            info["wav_in_output_directory"] = []

        # Try to read audio technical metadata via ffprobe (if available)
        try:
            result = subprocess.run(
                [
                    "ffprobe", "-v", "quiet",
                    "-print_format", "json",
                    "-show_format", "-show_streams",
                    str(audio_path),
                ],
                capture_output=True, text=True, timeout=15,
            )
            if result.returncode == 0:
                probe = json.loads(result.stdout)
                fmt = probe.get("format", {})
                streams = probe.get("streams", [])
                audio_streams = [s for s in streams if s.get("codec_type") == "audio"]
                tech: dict = {
                    "format_name": fmt.get("format_long_name", fmt.get("format_name")),
                    "duration_seconds": float(fmt.get("duration", 0)),
                    "duration_human": str(__import__("datetime").timedelta(seconds=int(float(fmt.get("duration", 0))))),
                    "bit_rate_kbps": round(int(fmt.get("bit_rate", 0)) / 1000, 1),
                    "size_bytes": int(fmt.get("size", 0)),
                }
                if audio_streams:
                    s = audio_streams[0]
                    tech["codec"] = s.get("codec_long_name", s.get("codec_name"))
                    tech["sample_rate_hz"] = int(s.get("sample_rate", 0))
                    tech["channels"] = s.get("channels")
                    tech["channel_layout"] = s.get("channel_layout")
                    tech["bits_per_sample"] = s.get("bits_per_sample")
                info["audio_technical"] = tech
        except Exception as exc:
            info["audio_technical"] = {"error": f"ffprobe unavailable or failed: {exc}"}

        # Check for merge provenance sidecar written by Function 0.
        # The sidecar lives next to the *original* merged file in the input directory
        # (named after the merged stem, e.g. "Kerry Bart (merged).merge_info.json").
        # audio_path is the output-directory copy (dg_<epoch>.mp3) so its stem never
        # matches — we must also check next to selected_path which retains the
        # original merged filename.
        merge_sidecar_candidates = [
            audio_path.parent / f"{audio_path.stem}.merge_info.json",      # output dir copy
            selected_path.parent / f"{selected_path.stem}.merge_info.json", # original input file
        ]
        for merge_sidecar in merge_sidecar_candidates:
            if merge_sidecar.exists():
                try:
                    with open(merge_sidecar, "r", encoding="utf-8") as _mf:
                        info["merge_info"] = json.load(_mf)
                except Exception:
                    pass
                break

        return info

    def execute_selected_function(function_key):
        """Execute the selected function from dropdown or show help if help mode is enabled"""
        if function_key and function_key in functions:
            # Check if help mode is enabled
            if help_mode_enabled.current and help_mode_enabled.current.value:
                # Show help dialog instead of executing
                show_help_dialog(function_key)
                # Clear selection
                active_function_dropdown.value = None
                page.update()
            else:
                # Execute the function normally
                # Call the function handler with a mock event
                class MockEvent:
                    pass

                functions[function_key]["handler"](MockEvent())

                # Clear selection after execution
                active_function_dropdown.value = None
                page.update()

    def get_sorted_function_options(function_list):
        """Get function dropdown options sorted by numeric order (0-5)"""
        # Extract function number from key (e.g., "function_0_merge_audio" -> 0)
        def get_function_number(func_key):
            import re
            match = re.search(r'function_(\d+)_', func_key)
            return int(match.group(1)) if match else 999

        # Sort by numeric order
        sorted_functions = sorted(function_list, key=get_function_number)

        # Create dropdown options
        options = []
        for func_key in sorted_functions:
            func_info = functions[func_key]
            label = f"{func_info['icon']} {func_info['label']}"
            options.append(ft.dropdown.Option(key=func_key, text=label))

        return options

    # ---- Collapsible directory controls
    dirs_expanded = True

    dirs_body_column = ft.Column(
        [
            ft.Text("Input Directory", size=14, weight=ft.FontWeight.W_500),
            ft.Row(
                [
                    input_directory_field,
                    ft.ElevatedButton(
                        "Browse...",
                        icon=ft.Icons.FOLDER_OPEN,
                        on_click=on_pick_directory_click,
                    ),
                    ft.ElevatedButton(
                        "Rescan",
                        icon=ft.Icons.REFRESH,
                        on_click=lambda e: _scan_audio_files(),
                        tooltip="Re-scan the input directory for WAV/MP3 files",
                    ),
                ],
                spacing=10,
            ),
            ft.Text("Working/Output Directory", size=14, weight=ft.FontWeight.W_500),
            ft.Column(
                [
                    ft.Row(
                        [
                            output_directory_field,
                            ft.ElevatedButton(
                                "Browse...",
                                icon=ft.Icons.FOLDER_OPEN,
                                on_click=on_pick_output_directory_click,
                            ),
                        ],
                        spacing=10,
                    ),
                    ft.Text(
                        "An 'OHM-data' subfolder will be created inside the selected working/output directory.",
                        size=11,
                        italic=True,
                        color=ft.Colors.GREY_600,
                    ),
                ],
                spacing=2,
            ),
            ft.Text("Select Permission PDF", size=14, weight=ft.FontWeight.W_500),
            ft.Row(
                [
                    pdf_selection_field,
                    ft.ElevatedButton(
                        "Pick...",
                        icon=ft.Icons.PICTURE_AS_PDF,
                        on_click=on_pick_pdf_click,
                        tooltip="Pick a permission/consent PDF from the input directory",
                    ),
                ],
                spacing=10,
            ),
        ],
        spacing=8,
    )

    inputs_header_row = ft.Row(
        [
            ft.Text("Inputs", size=18, weight=ft.FontWeight.BOLD),
        ],
        spacing=4,
        vertical_alignment=ft.CrossAxisAlignment.CENTER,
    )

    inputs_inner_column = ft.Column(
        [inputs_header_row, dirs_body_column],
        spacing=8,
    )

    def on_toggle_dirs(e):
        nonlocal dirs_expanded
        dirs_expanded = not dirs_expanded
        if dirs_expanded:
            inputs_inner_column.controls = [inputs_header_row, dirs_body_column]
        else:
            inputs_inner_column.controls = [inputs_header_row]
        dirs_toggle_button.icon = (
            ft.Icons.EXPAND_LESS if dirs_expanded else ft.Icons.EXPAND_MORE
        )
        page.update()

    dirs_toggle_button = ft.IconButton(
        icon=ft.Icons.EXPAND_LESS,
        icon_size=20,
        tooltip="Show/hide directory controls",
        on_click=on_toggle_dirs,
    )
    inputs_header_row.controls.append(dirs_toggle_button)

    # ------------------------------------------------------------------ layout

    page.add(
        ft.Column(
            controls=[
                # ---- Title
                ft.Text(
                    "🎙️ OHM — Oral History Manager",
                    size=24,
                    weight=ft.FontWeight.BOLD,
                ),
                ft.Text(
                    "A tool for managing creation and ingest of Oral Histories for Digital.Grinnell",
                    size=13,
                    color=ft.Colors.GREY_700,
                    italic=True,
                ),
                ft.Divider(height=5),

                # ---- Inputs section
                ft.Container(
                    content=inputs_inner_column,
                    padding=5,
                ),

                # ---- Audio file selection (fixed position, always visible)
                ft.Container(
                    content=ft.Row(
                        [
                            file_selection_field,
                            ft.ElevatedButton(
                                "Pick...",
                                icon=ft.Icons.AUDIO_FILE,
                                on_click=on_pick_file_click,
                                tooltip="Pick an audio file from the scanned list",
                            ),
                        ],
                        spacing=10,
                    ),
                    padding=ft.padding.only(left=5, right=5, bottom=5),
                ),

                ft.Divider(height=5),

                # ---- Functions section
                ft.Container(
                    content=ft.Column(
                        [
                            ft.Row(
                                [
                                    ft.Column(
                                        [
                                            ft.Text(
                                                "Active Functions",
                                                size=18,
                                                weight=ft.FontWeight.BOLD,
                                            ),
                                            ft.Text(
                                                "Select and execute workflow functions",
                                                size=12,
                                                italic=True,
                                                color=ft.Colors.GREY_700,
                                            ),
                                            ft.Container(height=5),
                                            active_function_dropdown := ft.Dropdown(
                                                label="Select Function to Execute",
                                                hint_text="Functions in workflow order (0-5)",
                                                width=500,
                                                options=[],
                                                on_change=lambda e: execute_selected_function(
                                                    e.control.value
                                                ),
                                            ),
                                            ft.Container(height=5),
                                            ft.Checkbox(
                                                label="Help Mode",
                                                ref=help_mode_enabled,
                                                tooltip="Enable to view help documentation for functions instead of executing them",
                                            ),
                                        ],
                                        spacing=5,
                                    ),
                                ],
                                vertical_alignment=ft.CrossAxisAlignment.START,
                            ),
                        ],
                        spacing=5,
                    ),
                    padding=5,
                ),

                ft.Divider(height=5),

                # ---- Status
                ft.Container(
                    content=ft.Column(
                        controls=[
                            ft.Row(
                                controls=[
                                    ft.Text(
                                        "Status",
                                        size=18,
                                        weight=ft.FontWeight.BOLD,
                                    ),
                                    ft.IconButton(
                                        icon=ft.Icons.COPY,
                                        tooltip="Copy status to clipboard",
                                        on_click=on_copy_status_click,
                                        icon_size=20,
                                    ),
                                ],
                            ),
                            status_container,
                        ],
                        spacing=5,
                    ),
                    padding=5,
                ),

                ft.Divider(height=5),

                # ---- Log output
                ft.Container(
                    content=ft.Column(
                        controls=[
                            ft.Row(
                                controls=[
                                    ft.Text(
                                        "Log Output",
                                        size=18,
                                        weight=ft.FontWeight.BOLD,
                                    ),
                                    ft.IconButton(
                                        icon=ft.Icons.DELETE_SWEEP,
                                        tooltip="Clear log",
                                        on_click=on_clear_log_click,
                                        icon_size=20,
                                    ),
                                ],
                            ),
                            ft.Container(
                                content=log_output,
                                border=ft.border.all(1, ft.Colors.GREY_400),
                                border_radius=5,
                                bgcolor=ft.Colors.GREY_100,
                            ),
                        ],
                        spacing=5,
                    ),
                    padding=5,
                ),
            ],
            spacing=4,
        )
    )

    # Populate function dropdowns with sorted options
    active_function_dropdown.options = get_sorted_function_options(active_functions)
    page.update()

    # Auto-scan on startup if a directory was restored from persistence
    if current_directory and current_directory.exists():
        _scan_audio_files()
        _scan_pdf_files()

    logger.info("UI initialised successfully")
    add_log_message("OHM application ready. Select a function to begin.")


if __name__ == "__main__":
    logger.info("Application starting…")
    ft.app(
        target=main,
        assets_dir="assets",
    )
